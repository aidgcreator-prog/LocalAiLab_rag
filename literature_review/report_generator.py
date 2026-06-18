"""Generate an APA 7th-edition Word document (.docx) literature review report.

Structure:
  Title page
  Abstract
  1. Search Methodology (query, APIs, n papers, n downloaded, n abstract-only)
  2. Findings  (populated from synthesis_text — agent-generated)
  3. Consensus and Convergence
  4. Debates and Divergence
  5. Research Gaps
  6. Conclusion
  References (APA 7th, sorted alphabetically, DOIs as hyperlinks)
  Appendix A — Search Log
"""

from __future__ import annotations

import datetime
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def synthesize_sections_from_rag(
    project: str,
    query: str,
    paper_records: list[dict[str, Any]],
    model_name: str = "",
    on_progress: Any = None,
) -> dict[str, str]:
    """Query the RAG vector store (literature review preset) and use the LLM to write
    each report section grounded exclusively in retrieved chunks.

    Uses the '\U0001f52c Literature review' preset:
        mode='Top-K Globally', top_k=12, fetch_k=250, max_files=8

    Returns dict with keys: findings, consensus, debates, gaps, conclusion.
    Falls back to empty string per section if RAG is unavailable.
    """
    from ragsub_agent.tools import rag_retrieve

    # Preset settings identical to the UI '\U0001f52c Literature review' preset
    _PRESET = dict(mode="Top-K Globally", top_k=12, fetch_k=250, max_files=8, min_rerank=0.02)

    # Build an APA citation map: title (lower) -> citation string
    cite_map: dict[str, str] = {
        (p.get("title") or "").lower(): _build_inline_citation(p)
        for p in paper_records
        if p.get("title")
    }

    # Per-section retrieval queries
    section_queries: dict[str, tuple[str, str]] = {
        "findings": (
            "findings",
            f"Key findings, results, and contributions in research about: {query}",
        ),
        "consensus": (
            "consensus",
            f"Points of agreement, convergence, and consensus across studies on: {query}",
        ),
        "debates": (
            "debates",
            f"Debates, disagreements, conflicting evidence, and controversies in: {query}",
        ),
        "gaps": (
            "gaps",
            f"Research gaps, limitations, future directions, and open questions in: {query}",
        ),
        "conclusion": (
            "conclusion",
            f"Conclusions, implications, and synthesis of evidence on: {query}",
        ),
    }

    _NO_DATA_SENTINEL = "NO_RELEVANT_DATA"

    # System prompt for each section — strict grounding, no hallucination
    _SECTION_INSTRUCTIONS: dict[str, str] = {
        "findings": (
            "You are writing the Findings section of an academic literature review. "
            "Use ONLY the retrieved context provided. Do NOT add any information, "
            "statistics, claims, or citations that are not explicitly present in the context. "
            "Do NOT invent authors, dates, or findings. "
            "If the retrieved context does not contain relevant information for this section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}"
        ),
        "consensus": (
            "You are writing the Consensus and Convergence section of an academic literature review. "
            "Describe only themes and findings explicitly supported by the retrieved context. "
            "Do NOT generalise beyond what the context states. "
            "Do NOT invent citations or agreement between papers. "
            "If the retrieved context does not contain relevant information for this section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}"
        ),
        "debates": (
            "You are writing the Debates and Divergence section of an academic literature review. "
            "Describe only disagreements or contradictions that are explicitly visible in the context. "
            "Do NOT infer or invent conflict between papers. "
            "If the retrieved context does not contain relevant information for this section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}"
        ),
        "gaps": (
            "You are writing the Research Gaps section of an academic literature review. "
            "Identify only gaps or limitations that are explicitly stated in the retrieved context. "
            "Do NOT speculate about gaps not mentioned in the context. "
            "If the retrieved context does not contain relevant information for this section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}"
        ),
        "conclusion": (
            "You are writing the Conclusion section of an academic literature review. "
            "Synthesise only what is explicitly stated in the retrieved context. "
            "Do NOT add external knowledge, general statements, or invented citations. "
            "If the retrieved context does not contain relevant information for this section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}"
        ),
    }

    sections: dict[str, str] = {}
    total = len(section_queries)

    # Resolve model
    _model = model_name or ""
    _llm = None
    if _model:
        try:
            from model_config import create_chat_model
            from langchain_core.messages import HumanMessage, SystemMessage
            _llm = create_chat_model(model_name=_model, temperature=0)
        except Exception:
            _llm = None

    for i, (key, (label, section_query)) in enumerate(section_queries.items()):
        if on_progress:
            on_progress("synthesis", i + 1, total, f"Querying RAG for: {label}…")

        # Retrieve chunks using the literature preset
        # rag_retrieve is a LangChain @tool (StructuredTool) — call via .invoke()
        context = rag_retrieve.invoke({
            "query": section_query,
            "project": project,
            "mode": _PRESET["mode"],
            "top_k": _PRESET["top_k"],
            "fetch_k": _PRESET["fetch_k"],
            "max_files": _PRESET["max_files"],
        })

        _rag_empty = (
            not context.strip()
            or context.strip().startswith("[WARN]")
        )

        if _rag_empty:
            sections[key] = ""
            continue

        if on_progress:
            on_progress("synthesis", i + 1, total, f"Writing {label} with LLM…")

        if _llm is None:
            # No LLM configured — use raw retrieved context with citations injected
            sections[key] = _inject_citations_in_text(context, paper_records)
            continue

        # Build LLM prompt with retrieved context
        from langchain_core.messages import HumanMessage, SystemMessage
        system_msg = SystemMessage(content=_SECTION_INSTRUCTIONS[key])
        human_msg = HumanMessage(content=(
            f"Retrieved context from the RAG database (project: '{project}'):\n"
            f"{'=' * 60}\n"
            f"{context}\n"
            f"{'=' * 60}\n\n"
            f"STRICT RULES:\n"
            f"1. Use ONLY the information present in the retrieved context above.\n"
            f"2. Do NOT add any external knowledge, general background, or invented citations.\n"
            f"3. Use APA in-text citations (e.g. (Smith et al., 2023)) that match sources in the context.\n"
            f"4. If the context does not contain relevant evidence for the '{label}' section, "
            f"respond with exactly: {_NO_DATA_SENTINEL}\n"
            f"5. Write 3-5 paragraphs. Do not add a section heading."
        ))
        try:
            response = _llm.invoke([system_msg, human_msg])
            text = (response.content if hasattr(response, "content") else str(response)).strip()
            if text == _NO_DATA_SENTINEL or text.startswith(_NO_DATA_SENTINEL):
                sections[key] = ""
            else:
                # Post-process: inject any missing APA citations
                sections[key] = _inject_citations_in_text(text, paper_records)
        except Exception as e:
            sections[key] = f"[Synthesis error: {e}]\n\n" + _inject_citations_in_text(
                context, paper_records
            )

    return sections


def generate_docx_report(
    project: str,
    query: str,
    synthesis_text: str,
    paper_records: list[dict[str, Any]],
    output_path: Path,
    *,
    pipeline_stats: dict[str, Any] | None = None,
    model_name: str = "",
    on_progress: Any = None,
) -> Path:
    """Generate a Word .docx literature review report.

    Args:
        project: Project/topic name used as the document title.
        query: Original search query string.
        synthesis_text: Agent-generated synthesis in Markdown-style sections.
            Expected sections (optional, extracted by heading):
            ## Findings, ## Consensus, ## Debates, ## Gaps, ## Conclusion
        paper_records: List of dicts with keys: title, authors, year, doi, journal,
            abstract, citation_count, method.
        output_path: Full path to write the .docx file.
        pipeline_stats: Optional dict with papers_found, papers_downloaded, etc.

    Returns:
        The output_path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_styles(doc)

    # ── Title Page ──────────────────────────────────────────────────────────
    _add_title_page(doc, project, query)

    # ── Abstract ────────────────────────────────────────────────────────────
    doc.add_heading("Abstract", level=1)
    n = len(paper_records)
    n_dl = sum(1 for p in paper_records if p.get("method") not in ("abstract_fallback", ""))
    n_ab = sum(1 for p in paper_records if p.get("method") == "abstract_fallback")
    abstract_text = (
        f"This automated literature review was conducted on the query: \"{query}\". "
        f"A total of {n} papers were identified across Semantic Scholar and OpenAlex. "
        f"Full texts were retrieved for {n_dl} papers; {n_ab} papers were included "
        f"using abstract-only content. The review synthesises key findings, points of "
        f"consensus, areas of debate, and identified research gaps."
    )
    doc.add_paragraph(abstract_text)

    # ── 1. Search Methodology ───────────────────────────────────────────────
    doc.add_heading("1. Search Methodology", level=1)
    stats = pipeline_stats or {}
    methodology_text = (
        f"The literature search was executed automatically using two open-access "
        f"academic databases: Semantic Scholar Graph API and OpenAlex. "
        f"The search query used was: \"{query}\". "
        f"Papers were ranked by open-access availability, citation count, and recency. "
        f"Results were deduplicated by DOI before final selection."
    )
    doc.add_paragraph(methodology_text)

    stat_table = doc.add_table(rows=1, cols=2)
    stat_table.style = "Table Grid"
    hdr = stat_table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Count"
    rows_data = [
        ("Papers identified (post-deduplication)", str(stats.get("papers_found", n))),
        ("Full-text PDFs downloaded", str(stats.get("papers_downloaded", n_dl))),
        ("Abstract-only inclusions", str(stats.get("papers_abstract_only", n_ab))),
        ("Papers excluded (no abstract/PDF)", str(stats.get("papers_skipped", 0))),
        ("Total RAG chunks ingested", str(stats.get("ingested_chunks", "—"))),
    ]
    for label, value in rows_data:
        row = stat_table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    # ── 2-6. Synthesis Sections ─────────────────────────────────────────────
    # Priority: 1) RAG+LLM synthesis  2) pasted synthesis_text  3) metadata fallback
    user_sections = _parse_synthesis_sections(synthesis_text)

    # Attempt RAG synthesis using the '🔬 Literature review' preset
    rag_sections: dict[str, str] = {}
    try:
        rag_sections = synthesize_sections_from_rag(
            project=project,
            query=query,
            paper_records=paper_records,
            model_name=model_name,
            on_progress=on_progress,
        )
    except Exception:
        rag_sections = {}

    section_map = {
        "findings": ("2. Findings", 2),
        "consensus": ("3. Consensus and Convergence", 3),
        "debates": ("4. Debates and Divergence", 4),
        "gaps": ("5. Research Gaps", 5),
        "conclusion": ("6. Conclusion", 6),
    }

    for key, (heading, num) in section_map.items():
        doc.add_heading(heading, level=1)

        rag_content = (rag_sections.get(key) or "").strip()
        user_content = (user_sections.get(key) or "").strip()
        # Suppress auto-metadata fallback to prevent hallucination.
        # Only use it when user explicitly pastes synthesis text or when
        # we have RAG content; otherwise say no data.
        auto_content = ""

        # Choose primary content: RAG first, then user paste
        if rag_content:
            # Tag so user knows it came from the RAG database
            _tag = doc.add_paragraph()
            _tag_run = _tag.add_run("\u2139\ufe0f Source: RAG database")
            _tag_run.italic = True
            _tag_run.font.size = Pt(9)
            _tag_run.font.color.rgb = RGBColor(0x44, 0x77, 0xAA)
            for para in rag_content.split("\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
        elif user_content:
            _tag = doc.add_paragraph()
            _tag_run = _tag.add_run("\u2139\ufe0f Source: pasted synthesis text")
            _tag_run.italic = True
            _tag_run.font.size = Pt(9)
            _tag_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            for para in user_content.split("\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
        else:
            _no_data_para = doc.add_paragraph()
            _no_data_run = _no_data_para.add_run(
                "No data available. No relevant content was found in the RAG "
                "database for this section. Please ingest documents and regenerate."
            )
            _no_data_run.italic = True
            _no_data_run.font.color.rgb = RGBColor(0xAA, 0x44, 0x44)

    # ── References ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("References", level=1)

    sorted_papers = sorted(paper_records, key=lambda p: _apa_sort_key(p))
    for paper in sorted_papers:
        _add_apa_reference(doc, paper)

    # ── Appendix A ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix A — Search Log", level=1)
    doc.add_paragraph(
        f"Search query: \"{query}\"\n"
        f"Search date: {datetime.date.today().isoformat()}\n"
        f"Databases: Semantic Scholar, OpenAlex\n"
        f"Project: {project}"
    )

    log_table = doc.add_table(rows=1, cols=5)
    log_table.style = "Table Grid"
    hdrs = log_table.rows[0].cells
    for i, h in enumerate(["#", "Title", "Authors", "Year", "Acquisition"]):
        hdrs[i].text = h

    for i, paper in enumerate(paper_records, 1):
        row = log_table.add_row().cells
        row[0].text = str(i)
        row[1].text = (paper.get("title") or "")[:80]
        authors = paper.get("authors") or []
        row[2].text = "; ".join(a.split(",")[0] for a in authors[:3])
        row[3].text = paper.get("year", "")
        row[4].text = paper.get("method", "unknown")

    doc.save(str(output_path))
    return output_path


# ── Helpers ───────────────────────────────────────────────────────────────────

def _configure_styles(doc: Document) -> None:
    """Apply consistent body text style."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)


def _add_title_page(doc: Document, project: str, query: str) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project)
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Automated Literature Review Report")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")

    doc.add_paragraph()

    query_para = doc.add_paragraph()
    query_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qrun = query_para.add_run(f"Search Query: \"{query}\"")
    qrun.italic = True

    doc.add_page_break()


def _parse_synthesis_sections(synthesis_text: str) -> dict[str, str]:
    """Extract named sections from agent synthesis Markdown output."""
    # Map heading keywords to section keys
    heading_map = {
        r"findings?|results?|overview": "findings",
        r"consensus|converge|agree|align": "consensus",
        r"debate|diverge|disagree|controver|conflict": "debates",
        r"gap|future|missing|lack|further": "gaps",
        r"conclusion|summary|synthesist": "conclusion",
    }

    sections: dict[str, str] = {}
    if not synthesis_text:
        return sections

    # Split by ## headings
    parts = re.split(r"\n##\s+", synthesis_text)
    for part in parts:
        if not part.strip():
            continue
        first_line, _, body = part.partition("\n")
        first_line = first_line.strip().lower()
        for pattern, key in heading_map.items():
            if re.search(pattern, first_line, re.I):
                sections[key] = body.strip()
                break
        else:
            # If no heading matched, put in findings as general content
            if "findings" not in sections:
                sections["findings"] = part.strip()

    # Also handle bare text (no ##) as findings
    if not sections and synthesis_text.strip():
        sections["findings"] = synthesis_text.strip()

    return sections


def _apa_sort_key(paper: dict[str, Any]) -> str:
    """Return sort key: first author surname + year."""
    authors = paper.get("authors") or []
    if authors:
        first = (authors[0] or "").split(",")[0].split()[-1] if authors[0] else "ZZZ"
    else:
        first = "ZZZ"
    return f"{first.lower()}_{paper.get('year', '0000')}"


def _add_apa_reference(doc: Document, paper: dict[str, Any]) -> None:
    """Add a single APA 7th-edition reference paragraph with hanging indent."""
    authors = paper.get("authors") or []
    year = paper.get("year", "n.d.")
    title = paper.get("title", "Untitled")
    journal = paper.get("journal", "")
    doi = paper.get("doi", "")

    author_str = _format_apa_authors(authors)
    ref = f"{author_str} ({year}). {title}."
    if journal:
        ref += f" {journal}."
    if doi:
        ref += f" https://doi.org/{doi}"

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.5)
    run = para.add_run(ref)
    run.font.size = Pt(11)


def _format_apa_authors(authors: list[str]) -> str:
    """Format author list in APA 7th style."""
    if not authors:
        return "Unknown Author."
    formatted = []
    for a in authors[:20]:
        a = a.strip()
        if not a:
            continue
        if "," in a:
            formatted.append(a)
        else:
            parts = a.split()
            if len(parts) >= 2:
                surname = parts[-1]
                initials = ". ".join(p[0] for p in parts[:-1] if p) + "."
                formatted.append(f"{surname}, {initials}")
            else:
                formatted.append(a)
    if len(authors) > 20:
        return "; ".join(formatted) + "; et al."
    return "; ".join(formatted) + ("." if formatted else "")


# ── In-text citation helpers ──────────────────────────────────────────────────

def _build_inline_citation(paper: dict[str, Any]) -> str:
    """Build APA 7th in-text citation, e.g. (Smith et al., 2024)."""
    authors = paper.get("authors") or []
    year = paper.get("year", "n.d.")
    if not authors:
        short_title = (paper.get("title") or "Untitled").split(":")[0][:30]
        return f'("{short_title}," {year})'
    first = (authors[0] or "").strip()
    if "," in first:
        surname = first.split(",")[0].strip()
    else:
        surname = first.split()[-1] if first.split() else first
    if len(authors) == 1:
        return f"({surname}, {year})"
    elif len(authors) == 2:
        second = (authors[1] or "").strip()
        if "," in second:
            surname2 = second.split(",")[0].strip()
        else:
            surname2 = second.split()[-1] if second.split() else second
        return f"({surname} & {surname2}, {year})"
    else:
        return f"({surname} et al., {year})"


def _first_author_surname(paper: dict[str, Any]) -> str:
    """Extract first-author surname for matching in text."""
    authors = paper.get("authors") or []
    if not authors:
        return ""
    first = (authors[0] or "").strip()
    if "," in first:
        return first.split(",")[0].strip()
    return first.split()[-1] if first.split() else ""


def _inject_citations_in_text(
    text: str,
    paper_records: list[dict[str, Any]],
) -> str:
    """Post-process text to append APA in-text citations where papers are
    mentioned by title or first-author surname without a nearby citation."""
    if not text or not paper_records:
        return text

    for paper in paper_records:
        cite = _build_inline_citation(paper)
        year = str(paper.get("year", ""))
        title = (paper.get("title") or "").strip()
        surname = _first_author_surname(paper)

        # Skip if citation is already present verbatim
        if cite in text:
            continue

        # Try matching title (first 40 chars, case-insensitive)
        if title and len(title) > 10:
            short_title = title[:40]
            idx = text.lower().find(short_title.lower())
            if idx != -1:
                # Check 60 chars after the title mention for an existing (Year)
                after_window = text[idx:idx + len(short_title) + 60]
                if year and re.search(r'\(' + re.escape(year) + r'\)', after_window):
                    continue  # citation already nearby
                end_pos = idx + len(short_title)
                # Find the end of the sentence or clause to insert
                while end_pos < len(text) and text[end_pos] not in ".;,\n":
                    end_pos += 1
                text = text[:end_pos] + " " + cite + text[end_pos:]
                continue  # one injection per paper

        # Try matching first-author surname (whole word, not already cited)
        if surname and len(surname) > 2:
            pattern = r'\b' + re.escape(surname) + r'\b'
            m = re.search(pattern, text)
            if m:
                after_window = text[m.start():m.start() + 80]
                if year and re.search(r'\(' + re.escape(year) + r'\)', after_window):
                    continue
                insert_at = m.end()
                text = text[:insert_at] + " " + cite + text[insert_at:]

    return text


# ── Auto-generated section content ────────────────────────────────────────────

def _auto_generate_sections(
    paper_records: list[dict[str, Any]],
) -> dict[str, str]:
    """Build structured section content from paper metadata & abstracts.

    Returns a dict with keys matching _parse_synthesis_sections output:
    findings, consensus, debates, gaps, conclusion.
    All text includes APA in-text citations.
    """
    if not paper_records:
        return {}

    cited = [
        (p, _build_inline_citation(p))
        for p in paper_records
        if p.get("title")
    ]
    years = [str(p.get("year", "")) for p in paper_records if p.get("year")]
    year_range = f"{min(years)}\u2013{max(years)}" if years else "various years"

    # ── Findings ──────────────────────────────────────────────────────────
    findings_paras: list[str] = []
    findings_paras.append(
        f"This review synthesizes {len(paper_records)} papers published "
        f"across {year_range}. The following summarizes the key contributions "
        f"of each work."
    )
    for paper, cite in cited:
        abstract = (paper.get("abstract") or "").strip()
        if abstract:
            sentences = re.split(r'(?<=[.!?])\s+', abstract)
            excerpt = " ".join(sentences[:2])
            if len(excerpt) > 300:
                excerpt = excerpt[:297] + "\u2026"
            findings_paras.append(f"{cite} \u2014 {excerpt}")
        else:
            title = paper.get("title", "Untitled")
            findings_paras.append(f"{cite} \u2014 \"{title}.\"")

    # ── Consensus ─────────────────────────────────────────────────────────
    consensus_paras: list[str] = []
    if len(cited) >= 2:
        consensus_paras.append(
            f"Across the reviewed literature, several thematic threads emerge. "
            f"The {len(paper_records)} studies address the topic from complementary "
            f"perspectives."
        )
        # High-citation papers
        high_cite = [
            (p, c) for p, c in cited if (p.get("citation_count") or 0) > 10
        ]
        if high_cite:
            consensus_paras.append(
                "Among the most-cited works, "
                + ", ".join(c for _, c in high_cite[:5])
                + " represent foundational contributions that subsequent "
                  "studies build upon."
            )
        # Group by journal
        journals: dict[str, list[str]] = {}
        for p, c in cited:
            j = (p.get("journal") or "").strip()
            if j:
                journals.setdefault(j, []).append(c)
        for jname, jcites in sorted(
            journals.items(), key=lambda x: -len(x[1])
        )[:3]:
            if len(jcites) >= 2:
                consensus_paras.append(
                    f"Multiple studies published in *{jname}* "
                    f"({', '.join(jcites[:3])}) suggest convergence in "
                    f"research focus within this venue."
                )
    else:
        consensus_paras.append(
            "With a limited number of studies reviewed, consensus patterns "
            "require further investigation."
        )

    # ── Debates ───────────────────────────────────────────────────────────
    debates_paras: list[str] = [
        "The reviewed literature presents varying methodological approaches "
        "and scope, reflecting differences in research perspectives.",
    ]
    if years and len(set(years)) > 1:
        oldest, newest = min(years), max(years)
        old_cites = [c for p, c in cited if str(p.get("year")) == oldest]
        new_cites = [c for p, c in cited if str(p.get("year")) == newest]
        if old_cites and new_cites:
            debates_paras.append(
                f"Earlier work {old_cites[0]} and more recent studies "
                f"{new_cites[0]} may reflect evolving methodologies and "
                f"shifting research priorities over the {oldest}\u2013{newest} period."
            )

    # ── Gaps ──────────────────────────────────────────────────────────────
    gaps_paras: list[str] = [
        "Based on the reviewed papers, several potential research gaps "
        "can be identified:",
    ]
    abstract_only = [
        (p, c) for p, c in cited if p.get("method") == "abstract_fallback"
    ]
    if abstract_only:
        gaps_paras.append(
            f"Full-text analysis was not possible for {len(abstract_only)} "
            f"paper(s) ({', '.join(c for _, c in abstract_only[:3])}), "
            f"indicating potential access barriers that limit comprehensive review."
        )
    gaps_paras.append(
        "Future work should expand the search scope, include additional "
        "databases, and consider longitudinal analysis to track how the "
        "field evolves over time."
    )

    # ── Conclusion ────────────────────────────────────────────────────────
    conclusion_paras: list[str] = []
    conclusion_paras.append(
        f"This automated literature review examined {len(paper_records)} papers "
        f"spanning {year_range}. The analysis reveals a growing body of evidence, "
        f"with key contributions from "
        + ", ".join(c for _, c in cited[:3])
        + (", and others." if len(cited) > 3 else ".")
    )
    conclusion_paras.append(
        "While this automated review provides a structured overview, a human "
        "expert review is recommended to validate the synthesized findings and "
        "interpret nuanced methodological differences across studies."
    )

    return {
        "findings": "\n\n".join(findings_paras),
        "consensus": "\n\n".join(consensus_paras),
        "debates": "\n\n".join(debates_paras),
        "gaps": "\n\n".join(gaps_paras),
        "conclusion": "\n\n".join(conclusion_paras),
    }
