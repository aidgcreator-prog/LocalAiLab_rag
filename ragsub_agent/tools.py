"""RAG tools for ingestion, retrieval, and reranking."""

from __future__ import annotations

import base64
import os
import re
import hashlib
import mimetypes
import subprocess
import tempfile
import time
import threading
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Any

from langchain_core.tools import tool
from model_config import create_chat_model, resolve_provider_and_model

PROJECT_DIR = Path(__file__).parent.parent
RAG_CHROMA_DIR = PROJECT_DIR / "rag-chroma"
RAG_ASSET_DIR = PROJECT_DIR / "tmp" / "rag_assets"
AUDIO_FILE_SUFFIXES = {".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus"}
VIDEO_FILE_SUFFIXES = {".mp4", ".mov", ".mkv", ".avi", ".webm"}
IMAGE_FILE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}

# ── Query-level result cache ─────────────────────────────────────────────
# TTL-based cache keyed on (query, project, themes, mode, top_k, fetch_k).
# Avoids re-running the full pipeline for identical interactive queries.
_retrieve_cache: dict[str, tuple[float, str]] = {}
_RETRIEVE_CACHE_TTL = float(os.getenv("RAG_CACHE_TTL_SECONDS", "120"))
_RETRIEVE_CACHE_MAX = 64
_retrieve_cache_lock = threading.Lock()


def _cache_key(query: str, project: str, themes: str, mode: str, top_k: int, fetch_k: int) -> str:
    raw = f"{query}|{project}|{themes}|{mode}|{top_k}|{fetch_k}"
    return hashlib.sha256(raw.encode()).hexdigest()[:24]


def _cache_get(key: str) -> str | None:
    with _retrieve_cache_lock:
        entry = _retrieve_cache.get(key)
        if entry is None:
            return None
        ts, result = entry
        if time.monotonic() - ts > _RETRIEVE_CACHE_TTL:
            del _retrieve_cache[key]
            return None
        return result


def _cache_put(key: str, result: str) -> None:
    with _retrieve_cache_lock:
        # Evict oldest if at capacity
        if len(_retrieve_cache) >= _RETRIEVE_CACHE_MAX:
            oldest = min(_retrieve_cache, key=lambda k: _retrieve_cache[k][0])
            del _retrieve_cache[oldest]
        _retrieve_cache[key] = (time.monotonic(), result)


# ── BM25 index cache ─────────────────────────────────────────────────────
# Persistent per-collection BM25 index rebuilt only when collection changes.
_bm25_index_cache: dict[str, tuple[int, Any, list[str], list[dict], list[str]]] = {}
_bm25_cache_lock = threading.Lock()


def _tokenize_for_bm25(text: str) -> list[str]:
    """Proper BM25 tokenization with punctuation removal and stop-word filtering."""
    _STOP_WORDS = frozenset({
        "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
        "of", "with", "by", "from", "is", "it", "as", "was", "are", "were",
        "be", "been", "being", "have", "has", "had", "do", "does", "did",
        "will", "would", "could", "should", "may", "might", "shall", "can",
        "this", "that", "these", "those", "not", "no", "nor", "so", "if",
        "than", "too", "very", "just", "about", "into", "through", "during",
        "before", "after", "above", "below", "between", "each", "all", "both",
        "such", "its", "own", "same", "other", "which", "who", "whom",
        "what", "when", "where", "how", "there", "here", "also", "only",
    })
    tokens = re.findall(r"[a-zA-Z0-9]+", text.lower())
    return [t for t in tokens if len(t) > 1 and t not in _STOP_WORDS]


def _get_bm25_index(
    collection_name: str,
    collection_obj: Any,
    where_filter: dict | None,
    bm25_cap: int,
) -> tuple[Any, list[str], list[dict], list[str]]:
    """Get or build a cached BM25 index for a collection.

    Rebuilds when collection count changes (new ingestion).
    Returns (bm25_model, documents, metadatas, ids).
    """
    c_count = collection_obj.count()
    with _bm25_cache_lock:
        cached = _bm25_index_cache.get(collection_name)
        if cached and cached[0] == c_count:
            _, bm25_model, docs, metas, ids = cached
            # Apply where_filter client-side
            if where_filter:
                filtered = [
                    (d, m, i) for d, m, i in zip(docs, metas, ids)
                    if _meta_matches_filter(m, where_filter)
                ]
                if filtered:
                    return bm25_model, [t[0] for t in filtered], [t[1] for t in filtered], [t[2] for t in filtered]
                return bm25_model, [], [], []
            return bm25_model, docs, metas, ids

    # Cache miss — rebuild
    c_data = collection_obj.get(include=["documents", "metadatas"], limit=bm25_cap)
    docs = [str(d) for d in (c_data.get("documents") or [])]
    metas = [m or {} for m in (c_data.get("metadatas") or [])]
    ids = [str(i) for i in (c_data.get("ids") or [])]

    try:
        from rank_bm25 import BM25Okapi
        tokenized = [_tokenize_for_bm25(d) for d in docs]
        bm25_model = BM25Okapi(tokenized) if tokenized else None
    except ImportError:
        bm25_model = None

    with _bm25_cache_lock:
        _bm25_index_cache[collection_name] = (c_count, bm25_model, docs, metas, ids)

    if where_filter:
        filtered = [
            (d, m, i) for d, m, i in zip(docs, metas, ids)
            if _meta_matches_filter(m, where_filter)
        ]
        if filtered:
            return bm25_model, [t[0] for t in filtered], [t[1] for t in filtered], [t[2] for t in filtered]
        return bm25_model, [], [], []
    return bm25_model, docs, metas, ids


# ── HyDE (Hypothetical Document Embeddings) ──────────────────────────────

def _generate_hyde_document(query: str) -> str:
    """Generate a hypothetical document that would answer the query.

    The hypothetical answer is embedded instead of the raw query, which
    moves the query vector closer to where actual answers live in the
    embedding space. Improves zero-shot recall by 10-30%.
    Returns empty string on failure (caller falls back to raw query).
    """
    prompt = (
        "Write a short, factual paragraph (3-5 sentences) that would directly "
        "answer the following question. Write as if you are an expert author "
        "writing a textbook passage. Do not say 'I think' or 'This is about'. "
        "Just write the answer passage directly.\n\n"
        f"Question: {query}\n\n"
        "Answer passage:"
    )
    try:
        llm = _get_context_llm()
        response = llm.invoke(prompt)
        text = _extract_text_from_llm_response(response)
        return text[:1000] if text else ""
    except Exception:
        return ""


# ── Query Decomposition ──────────────────────────────────────────────────

def _decompose_query(query: str) -> list[str]:
    """Break a complex query into 2-3 focused sub-queries for better recall.

    Only decomposes when the query is complex (contains multiple aspects,
    conjunctions, or is longer than ~15 words). Returns [original_query]
    for simple queries.
    """
    # Heuristic: only decompose if query is complex enough
    word_count = len(query.split())
    has_conjunction = bool(re.search(r'\b(and|or|versus|vs|compared|between|relationship|both|also|as well as)\b', query, re.I))
    has_multiple_questions = query.count("?") > 1
    if word_count < 10 and not has_conjunction and not has_multiple_questions:
        return [query]

    prompt = (
        "Break this research query into 2-3 simpler, focused sub-queries that "
        "together cover all aspects of the original question. Each sub-query "
        "should target a different aspect. Return ONLY the sub-queries, one per "
        "line, no numbering or bullets.\n\n"
        f"Query: {query}\n\n"
        "Sub-queries:"
    )
    try:
        llm = _get_context_llm()
        response = llm.invoke(prompt)
        text = _extract_text_from_llm_response(response)
        sub_queries = [q.strip().lstrip("0123456789.-) ") for q in text.strip().split("\n") if q.strip()]
        # Filter out empty or too-short sub-queries
        sub_queries = [q for q in sub_queries if len(q) > 5]
        if not sub_queries:
            return [query]
        # Always include the original query for full-context matching
        return [query] + sub_queries[:3]
    except Exception:
        return [query]


def _get_rag_collection_name(project: str = "Default") -> str:
    """Return a stable Chroma collection name scoped to project and embed model.

    Format: rag__{project}__{model_hash}
    - project is normalised to alphanumeric+underscore, max 20 chars.
    - model_hash is the first 8 hex chars of SHA-1(embed_model_name).
    """
    embed_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
    model_hash = hashlib.sha1(embed_model.encode("utf-8")).hexdigest()[:8]
    proj_normalized = re.sub(r"[^a-zA-Z0-9]+", "_", project).strip("_").lower()[:20]
    return f"rag__{proj_normalized}__{model_hash}"


class _LlamaCppEmbeddingFunction:
    """Chroma-compatible embedding function backed by llama-cpp-python."""

    def __init__(
        self,
        model_path: str,
        *,
        n_ctx: int,
        n_batch: int,
        n_gpu_layers: int,
        flash_attn: bool,
    ) -> None:
        self.model_path = model_path
        self.n_ctx = n_ctx
        self.n_batch = n_batch
        self.n_gpu_layers = n_gpu_layers
        self.flash_attn = flash_attn
        self._model: Any | None = None

    def _get_model(self):
        if self._model is not None:
            return self._model

        try:
            from llama_cpp import Llama
        except Exception as exc:
            raise RuntimeError(
                "llama.cpp embeddings require the `llama-cpp-python` package. "
                "Install the llama_cpp extra and retry."
            ) from exc

        self._model = Llama(
            model_path=self.model_path,
            embedding=True,
            n_ctx=self.n_ctx,
            n_batch=self.n_batch,
            n_gpu_layers=self.n_gpu_layers,
            flash_attn=self.flash_attn,
            verbose=False,
        )
        return self._model

    def __call__(self, input: list[str]) -> list[list[float]]:
        texts = [str(item) for item in input]
        try:
            response = self._get_model().create_embedding(texts)
        except Exception as exc:
            raise RuntimeError(
                f"Failed to generate llama.cpp embeddings using {self.model_path}: {exc}"
            ) from exc

        embeddings = [item.get("embedding") for item in (response.get("data") or [])]
        if len(embeddings) != len(texts) or any(emb is None for emb in embeddings):
            raise RuntimeError(
                f"llama.cpp returned an unexpected embedding payload for {self.model_path}."
            )
        return [list(embedding) for embedding in embeddings if embedding is not None]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.__call__(texts)

    def embed_query(self, text: str) -> list[float]:
        return self.__call__([text])[0]


class _LangChainEmbeddingAdapter:
    """Chroma-compatible wrapper around LangChain embedding models."""

    def __init__(self, embeddings: Any) -> None:
        self.embeddings = embeddings

    def __call__(self, input: list[str]) -> list[list[float]]:
        texts = [str(item) for item in input]
        try:
            return self.embeddings.embed_documents(texts)
        except Exception as exc:
            raise RuntimeError(f"Failed to generate embeddings: {exc}") from exc

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.__call__(texts)

    def embed_query(self, text: str) -> list[float]:
        try:
            return self.embeddings.embed_query(text)
        except Exception as exc:
            raise RuntimeError(f"Failed to generate query embedding: {exc}") from exc


@lru_cache(maxsize=8)
def _get_rag_embedding_function(
    embedding_model: str,
    llama_server_base_url: str,
    llama_cpp_n_ctx: int,
    llama_cpp_n_batch: int,
    llama_cpp_n_gpu_layers: int,
    llama_cpp_flash_attn: bool,
):
    """Return the configured embedding backend for RAG ingestion and retrieval."""
    provider, resolved_model = resolve_provider_and_model(embedding_model)

    if provider == "llama_cpp":
        if not resolved_model:
            raise ValueError(
                "llama.cpp embeddings require a GGUF model path. Set RAG_EMBED_MODEL "
                "to `llama_cpp:/absolute/path/to/model.gguf`."
            )
        return _LlamaCppEmbeddingFunction(
            resolved_model,
            n_ctx=llama_cpp_n_ctx,
            n_batch=llama_cpp_n_batch,
            n_gpu_layers=llama_cpp_n_gpu_layers,
            flash_attn=llama_cpp_flash_attn,
        )

    if provider == "llama_server":
        try:
            from langchain_openai import OpenAIEmbeddings
        except Exception as exc:
            raise RuntimeError(
                "llama-server embeddings require the `langchain-openai` package."
            ) from exc

        return _LangChainEmbeddingAdapter(
            OpenAIEmbeddings(
                base_url=llama_server_base_url,
                api_key="llama-server",
                model=resolved_model or os.getenv("LLAMA_SERVER_MODEL", "local").strip() or "local",
            )
        )

    try:
        from langchain_ollama import OllamaEmbeddings
    except Exception as exc:
        raise RuntimeError(
            "RAG embeddings require either llama.cpp support or the Ollama "
            "embedding wrapper."
        ) from exc

    return _LangChainEmbeddingAdapter(
        OllamaEmbeddings(
            model=resolved_model or embedding_model,
            base_url=os.getenv("RAG_OLLAMA_BASE_URL", "http://localhost:11434"),
        )
    )


class _LlamaCppVisionModel:
    """Cached llama.cpp multimodal captioning model."""

    def __init__(
        self,
        model_path: str,
        *,
        chat_format: str | None,
        n_ctx: int,
        n_batch: int,
        n_gpu_layers: int,
        flash_attn: bool,
    ) -> None:
        self.model_path = model_path
        self.chat_format = chat_format
        self.n_ctx = n_ctx
        self.n_batch = n_batch
        self.n_gpu_layers = n_gpu_layers
        self.flash_attn = flash_attn
        self._model: Any | None = None

    def _get_model(self):
        if self._model is not None:
            return self._model

        try:
            from llama_cpp import Llama
        except Exception as exc:
            raise RuntimeError(
                "llama.cpp vision captioning requires the `llama-cpp-python` package."
            ) from exc

        kwargs: dict[str, Any] = {
            "model_path": self.model_path,
            "n_ctx": self.n_ctx,
            "n_batch": self.n_batch,
            "n_gpu_layers": self.n_gpu_layers,
            "flash_attn": self.flash_attn,
            "verbose": False,
        }
        if self.chat_format:
            kwargs["chat_format"] = self.chat_format

        self._model = Llama(**kwargs)
        return self._model

    def caption(self, image_bytes: bytes, prompt: str) -> str:
        data_uri = "data:image/png;base64," + base64.b64encode(image_bytes).decode("ascii")
        messages = [
            {
                "role": "system",
                "content": "You write concise, factual image captions for retrieval.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_uri}},
                ],
            },
        ]
        response = self._get_model().create_chat_completion(
            messages=messages,
            temperature=0.2,
            max_tokens=256,
        )
        choice = (response.get("choices") or [{}])[0]
        message = choice.get("message") or {}
        content = message.get("content") if isinstance(message, dict) else getattr(message, "content", None)
        return _extract_text_from_content(content)


@lru_cache(maxsize=8)
def _get_rag_vision_model(
    vision_model: str,
    llama_server_base_url: str,
    llama_cpp_n_ctx: int,
    llama_cpp_n_batch: int,
    llama_cpp_n_gpu_layers: int,
    llama_cpp_flash_attn: bool,
    llama_cpp_chat_format: str,
):
    provider, resolved_model = resolve_provider_and_model(vision_model)

    if provider == "llama_cpp":
        if not resolved_model:
            raise ValueError(
                "llama.cpp vision captioning requires a GGUF model path. Set RAG_VISION_MODEL "
                "to `llama_cpp:/absolute/path/to/model.gguf`."
            )
        return _LlamaCppVisionModel(
            resolved_model,
            chat_format=llama_cpp_chat_format or None,
            n_ctx=llama_cpp_n_ctx,
            n_batch=llama_cpp_n_batch,
            n_gpu_layers=llama_cpp_n_gpu_layers,
            flash_attn=llama_cpp_flash_attn,
        )

    if provider == "llama_server":
        return ("llama_server", resolved_model or os.getenv("LLAMA_SERVER_MODEL", "local").strip() or "local")

    return ("ollama", resolved_model or vision_model)


def _normalize_file_list(file_paths: str) -> list[Path]:
    parts = [p.strip() for p in re.split(r"[\n,;]+", file_paths or "") if p.strip()]
    resolved: list[Path] = []
    for part in parts:
        path = Path(part)
        if not path.is_absolute():
            path = PROJECT_DIR / path
        resolved.append(path)
    return resolved


def _parse_themes(themes: str) -> list[str]:
    return [t.strip() for t in re.split(r"[,;]+", themes or "") if t.strip()]


def _parse_modalities(modalities: str) -> list[str]:
    return [
        label
        for label in (part.strip().lower() for part in re.split(r"[,;]+", modalities or ""))
        if label in {"text", "table", "image"}
    ]


def _normalize_project(project: str | None) -> str:
    """Return a stable default project name for RAG admin operations."""
    return (project or "").strip() or "Default"


def _normalize_id_part(value: str, max_len: int = 24) -> str:
    """Normalize a label for use inside stable chunk IDs."""
    normalized = re.sub(r"[^a-zA-Z0-9]+", "_", value or "").strip("_").lower()
    return (normalized or "na")[:max_len]


def _fingerprint_text(value: str, length: int = 12) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:length]


def _build_source_fingerprint(path: Path) -> str:
    """Fingerprint a source path so same filenames in different folders stay isolated."""
    return _fingerprint_text(str(path.resolve()))


def _build_chunk_fingerprint(
    text: str,
    *,
    index: int,
    section: str,
    page_number: str,
) -> str:
    fingerprint_input = f"{index}|{page_number}|{section}|{text.strip()}"
    return _fingerprint_text(fingerprint_input)


def _build_chunk_id(
    project: str,
    theme: str,
    source_fingerprint: str,
    chunk_fingerprint: str,
) -> str:
    return (
        f"{_normalize_id_part(project)}__{_normalize_id_part(theme)}__"
        f"{source_fingerprint}__{chunk_fingerprint}"
    )


def _safe_asset_relative_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(PROJECT_DIR.resolve()).as_posix()
    except Exception:
        return str(path)


def _normalize_text_for_overlap(text: str) -> set[str]:
    """Convert text into a lightweight token set for overlap-based dedup checks."""
    return {
        token
        for token in re.findall(r"[a-zA-Z0-9]+", (text or "").lower())
        if len(token) >= 3
    }


def _compute_token_overlap_ratio(left: str, right: str) -> float:
    """Return token-overlap ratio using the smaller token set as denominator."""
    left_tokens = _normalize_text_for_overlap(left)
    right_tokens = _normalize_text_for_overlap(right)
    if not left_tokens or not right_tokens:
        return 0.0
    shared = left_tokens & right_tokens
    return len(shared) / max(1, min(len(left_tokens), len(right_tokens)))


def _dedupe_multimodal_payloads(
    base_docs: list[str],
    base_metas: list[dict[str, str]],
    candidate_docs: list[str],
    candidate_metas: list[dict[str, str]],
    candidate_ids: list[str],
) -> tuple[list[str], list[dict[str, str]], list[str]]:
    """Drop multimodal chunks that substantially duplicate existing text chunks."""
    if not candidate_docs:
        return candidate_docs, candidate_metas, candidate_ids

    min_overlap = float(os.getenv("RAG_MULTIMODAL_DEDUP_OVERLAP", "0.85"))
    base_by_page: dict[tuple[str, str], list[str]] = {}
    for base_doc, base_meta in zip(base_docs, base_metas):
        base_modality = str((base_meta or {}).get("modality", "text") or "text").lower()
        if base_modality != "text":
            continue
        key = (
            str((base_meta or {}).get("source_fingerprint", "")),
            str((base_meta or {}).get("page_number", "")),
        )
        base_by_page.setdefault(key, []).append(base_doc)

    kept_docs: list[str] = []
    kept_metas: list[dict[str, str]] = []
    kept_ids: list[str] = []
    kept_candidate_docs_by_page: dict[tuple[str, str], list[str]] = {}

    for candidate_doc, candidate_meta, candidate_id in zip(candidate_docs, candidate_metas, candidate_ids):
        key = (
            str((candidate_meta or {}).get("source_fingerprint", "")),
            str((candidate_meta or {}).get("page_number", "")),
        )
        comparisons = list(base_by_page.get(key, [])) + kept_candidate_docs_by_page.get(key, [])
        duplicate = any(
            _compute_token_overlap_ratio(candidate_doc, existing_doc) >= min_overlap
            for existing_doc in comparisons
        )
        if duplicate:
            continue
        kept_docs.append(candidate_doc)
        kept_metas.append(candidate_meta)
        kept_ids.append(candidate_id)
        kept_candidate_docs_by_page.setdefault(key, []).append(candidate_doc)

    return kept_docs, kept_metas, kept_ids


@lru_cache(maxsize=1)
def _tesseract_available() -> bool:
    configured = os.getenv("TESSERACT_CMD", "").strip()
    if configured and Path(configured).exists():
        return True
    return shutil.which("tesseract") is not None


def _get_tesseract_ocr_languages() -> str:
    """Return the configured OCR language list for Tesseract."""
    raw = os.getenv("RAG_OCR_LANGS", "").strip()
    return raw or "eng"


def _get_ffmpeg_executable() -> str:
    """Return a usable ffmpeg executable path, preferring imageio-ffmpeg when installed."""
    resolved = shutil.which("ffmpeg")
    if resolved:
        return resolved
    try:
        import imageio_ffmpeg

        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception as exc:
        raise RuntimeError(
            "ffmpeg is required for local audio/video ingestion. Install ffmpeg or add imageio-ffmpeg."
        ) from exc


def get_multimodal_ocr_status() -> dict[str, str | bool]:
    """Return OCR runtime availability for multimodal RAG image enrichment."""
    try:
        import pytesseract  # noqa: F401

        pytesseract_ok = True
    except Exception:
        pytesseract_ok = False

    configured = os.getenv("TESSERACT_CMD", "").strip()
    if configured:
        cmd_ok = Path(configured).exists()
        cmd_label = configured
        configured_label = configured
        detected_label = ""
    else:
        resolved = shutil.which("tesseract")
        cmd_ok = resolved is not None
        cmd_label = resolved or ""
        configured_label = ""
        detected_label = resolved or ""

    remediation = (
        "Install Tesseract and set "
        "TESSERACT_CMD=C:\\Program Files\\Tesseract-OCR\\tesseract.exe "
        "or add tesseract to PATH."
    )
    ocr_langs = _get_tesseract_ocr_languages()

    if pytesseract_ok and cmd_ok:
        return {
            "enabled": True,
            "state": "ready",
            "message": f"OCR enabled via Tesseract ({cmd_label or 'PATH'}) using languages: {ocr_langs}",
            "configured_path": configured_label,
            "detected_path": detected_label or cmd_label,
            "remediation": "",
        }
    if pytesseract_ok and not cmd_ok:
        return {
            "enabled": False,
            "state": "missing_binary",
            "message": (
                "pytesseract installed, but the Tesseract binary is missing. "
                f"{remediation}"
            ),
            "configured_path": configured_label,
            "detected_path": detected_label,
            "remediation": remediation,
        }
    if not pytesseract_ok and cmd_ok:
        return {
            "enabled": False,
            "state": "missing_python",
            "message": "Tesseract binary found, but pytesseract is not installed in the active Python environment.",
            "configured_path": configured_label,
            "detected_path": detected_label or cmd_label,
            "remediation": "Install pytesseract in the active Python environment.",
        }
    return {
        "enabled": False,
        "state": "disabled",
        "message": "OCR unavailable: install pytesseract and the Tesseract binary to enrich multimodal image records.",
        "configured_path": configured_label,
        "detected_path": detected_label,
        "remediation": remediation,
    }


def _extract_image_ocr_text(image_bytes: bytes) -> str:
    """Best-effort OCR for extracted images using pytesseract when available."""
    if not _tesseract_available():
        return ""
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    configured = os.getenv("TESSERACT_CMD", "").strip()
    if configured:
        try:
            pytesseract.pytesseract.tesseract_cmd = configured
        except Exception:
            return ""

    try:
        with Image.open(BytesIO(image_bytes)) as image:
            text = pytesseract.image_to_string(image, lang=_get_tesseract_ocr_languages())
    except Exception:
        return ""
    return re.sub(r"\s+", " ", (text or "").strip())[:1200]


def _extract_audio_track_for_transcription(file_path: Path) -> Path:
    """Extract a mono MP3 audio track from a video file for transcription."""
    ffmpeg_exe = _get_ffmpeg_executable()
    temp_dir = Path(tempfile.mkdtemp(prefix="rag_media_audio_", dir=str(PROJECT_DIR / "tmp")))
    output_path = temp_dir / f"{file_path.stem}.mp3"
    command = [
        ffmpeg_exe,
        "-y",
        "-i",
        str(file_path),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        str(output_path),
    ]
    result = subprocess.run(command, capture_output=True, text=True)
    if result.returncode != 0 or not output_path.exists():
        raise RuntimeError(f"ffmpeg failed to extract audio from {file_path.name}: {result.stderr.strip()}")
    return output_path


def _get_local_transcribe_model():
    """Return a cached local faster-whisper model instance."""
    from faster_whisper import WhisperModel

    model_name = os.getenv("RAG_LOCAL_TRANSCRIBE_MODEL", "small").strip() or "small"
    cache = getattr(_get_local_transcribe_model, "_cache", {})
    if cache.get("model") != model_name:
        cache = {
            "model": model_name,
            "instance": WhisperModel(
                model_name,
                device=os.getenv("RAG_LOCAL_TRANSCRIBE_DEVICE", "cpu"),
                compute_type=os.getenv("RAG_LOCAL_TRANSCRIBE_COMPUTE_TYPE", "int8"),
            ),
        }
        _get_local_transcribe_model._cache = cache
    return cache["instance"]


def _transcribe_media_file_local(file_path: Path) -> str:
    """Transcribe media locally using faster-whisper."""
    model = _get_local_transcribe_model()
    segments, _info = model.transcribe(str(file_path), vad_filter=True)
    transcript = " ".join((segment.text or "").strip() for segment in segments).strip()
    if not transcript:
        raise RuntimeError(f"Local transcription returned empty text for {file_path.name}.")
    return transcript


def _transcribe_media_file(file_path: Path) -> str:
    """Transcribe a local audio or video file using OpenAI if configured, else local Whisper."""
    api_key = os.getenv("OPENAI_API_KEY", "").strip()

    working_path = file_path
    cleanup_path: Path | None = None
    if file_path.suffix.lower() in VIDEO_FILE_SUFFIXES:
        cleanup_path = _extract_audio_track_for_transcription(file_path)
        working_path = cleanup_path

    try:
        if api_key:
            api_base = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/")
            if not api_base.endswith("/v1"):
                api_base = f"{api_base}/v1"
            endpoint = f"{api_base}/audio/transcriptions"
            model = os.getenv("RAG_TRANSCRIBE_MODEL", os.getenv("VOICE_TRANSCRIBE_MODEL", "gpt-4o-mini-transcribe")).strip()
            mime_type = mimetypes.guess_type(working_path.name)[0] or "application/octet-stream"
            with working_path.open("rb") as handle:
                files = {"file": (working_path.name, handle, mime_type)}
                data = {"model": model}
                headers = {"Authorization": f"Bearer {api_key}"}
                with httpx.Client(timeout=120.0) as client:
                    response = client.post(endpoint, headers=headers, data=data, files=files)
                    response.raise_for_status()
            if "application/json" in response.headers.get("content-type", ""):
                payload = response.json()
                transcript = str(payload.get("text", "") or "").strip()
            else:
                transcript = response.text.strip()
        else:
            transcript = _transcribe_media_file_local(working_path)
        if not transcript:
            raise RuntimeError(f"Transcription returned empty text for {file_path.name}.")
        return transcript
    finally:
        if cleanup_path is not None:
            try:
                cleanup_path.unlink(missing_ok=True)
                cleanup_path.parent.rmdir()
            except Exception:
                pass


def get_multimodal_vision_status() -> dict[str, str | bool]:
    """Return runtime availability for optional vision captioning."""
    raw_model = os.getenv("RAG_VISION_MODEL", "").strip()
    if not raw_model:
        return {
            "enabled": False,
            "state": "disabled",
            "message": "Vision captioning disabled: set RAG_VISION_MODEL to a llama.cpp vision model.",
        }

    provider, model = resolve_provider_and_model(raw_model)

    if provider == "llama_cpp":
        if not model:
            return {
                "enabled": False,
                "state": "missing_model",
                "message": "Vision model configured, but no llama.cpp GGUF path was provided.",
            }
        if not Path(model).expanduser().exists():
            return {
                "enabled": False,
                "state": "missing_model",
                "message": f"llama.cpp vision model not found: {model}",
            }
        try:
            chat_format = os.getenv("RAG_VISION_CHAT_FORMAT", "").strip()
            llama_cpp_n_ctx = int(os.getenv("RAG_VISION_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
            llama_cpp_n_batch = int(os.getenv("RAG_VISION_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
            llama_cpp_n_gpu_layers = int(os.getenv("RAG_VISION_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
            llama_cpp_flash_attn = os.getenv(
                "RAG_VISION_FLASH_ATTN",
                os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
            ).strip().lower() in {"1", "true", "yes", "on"}
            _get_rag_vision_model(
                raw_model,
                "",
                llama_cpp_n_ctx,
                llama_cpp_n_batch,
                llama_cpp_n_gpu_layers,
                llama_cpp_flash_attn,
                chat_format,
            )
        except Exception as exc:
            return {
                "enabled": False,
                "state": "unreachable",
                "message": f"llama.cpp vision model could not be loaded: {exc}",
            }
        return {
            "enabled": True,
            "state": "ready",
            "message": f"Vision captioning enabled via llama.cpp model {Path(model).name}.",
        }

    if provider == "llama_server":
        base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
        try:
            import httpx

            with httpx.Client(timeout=10.0) as client:
                response = client.get(f"{base_url}/models")
                response.raise_for_status()
        except Exception as exc:
            return {
                "enabled": False,
                "state": "unreachable",
                "message": f"Vision model configured ({model}), but llama-server is unreachable at {base_url}.",
            }
        return {
            "enabled": True,
            "state": "ready",
            "message": f"Vision captioning enabled via llama-server model {model or 'local'}.",
        }

    try:
        import httpx

        base_url = os.getenv("RAG_OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
        with httpx.Client(timeout=10.0) as client:
            response = client.get(f"{base_url}/api/tags")
            response.raise_for_status()
            payload = response.json()
    except Exception:
        return {
            "enabled": False,
            "state": "unreachable",
            "message": f"Vision model configured ({model}), but Ollama is unreachable at {base_url}.",
        }

    model_names = {
        (item.get("name") or "").strip()
        for item in (payload.get("models") or [])
        if isinstance(item, dict)
    }
    if model in model_names or any(name.startswith(f"{model}:") for name in model_names):
        return {
            "enabled": True,
            "state": "ready",
            "message": f"Vision captioning enabled via Ollama model {model}.",
        }
    return {
        "enabled": False,
        "state": "missing_model",
        "message": f"Vision model {model} is not available in Ollama. Pull it first.",
    }


def _extract_image_vision_caption(image_bytes: bytes) -> str:
    """Best-effort image captioning using an optional vision model."""
    raw_model = os.getenv("RAG_VISION_MODEL", "").strip()
    if not raw_model:
        return ""

    provider, model = resolve_provider_and_model(raw_model)
    prompt = (
        "Describe the image for retrieval in 1-3 factual sentences. "
        "Mention visible objects, diagrams, labels, UI elements, charts, or screenshots. "
        "Do not speculate beyond what is visible."
    )

    try:
        if provider == "llama_cpp":
            chat_format = os.getenv("RAG_VISION_CHAT_FORMAT", "").strip()
            llama_cpp_n_ctx = int(os.getenv("RAG_VISION_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
            llama_cpp_n_batch = int(os.getenv("RAG_VISION_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
            llama_cpp_n_gpu_layers = int(os.getenv("RAG_VISION_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
            llama_cpp_flash_attn = os.getenv(
                "RAG_VISION_FLASH_ATTN",
                os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
            ).strip().lower() in {"1", "true", "yes", "on"}
            vision_model = _get_rag_vision_model(
                raw_model,
                "",
                llama_cpp_n_ctx,
                llama_cpp_n_batch,
                llama_cpp_n_gpu_layers,
                llama_cpp_flash_attn,
                chat_format,
            )
            return vision_model.caption(image_bytes, prompt)[:1200]

        if provider == "llama_server":
            base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
            try:
                from langchain_openai import ChatOpenAI
            except Exception:
                return ""
            llm = ChatOpenAI(
                base_url=base_url,
                api_key="llama-server",
                model=model or os.getenv("LLAMA_SERVER_MODEL", "local").strip() or "local",
                temperature=0.2,
            )
            data_uri = "data:image/png;base64," + base64.b64encode(image_bytes).decode("ascii")
            response = llm.invoke([
                {
                    "role": "system",
                    "content": "You write concise, factual image captions for retrieval.",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_uri}},
                    ],
                },
            ])
            return _extract_text_from_llm_response(response)[:1200]

        import httpx

        base_url = os.getenv("RAG_OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "images": [base64.b64encode(image_bytes).decode("ascii")],
        }
        with httpx.Client(timeout=30.0) as client:
            response = client.post(f"{base_url}/api/generate", json=payload)
            response.raise_for_status()
            data = response.json()
    except Exception:
        return ""

    caption = re.sub(r"\s+", " ", str(data.get("response", "") or "").strip())
    return caption[:1200]


def _should_extract_structured_visual_notes(
    page_text: str,
    ocr_text: str,
    vision_caption: str,
    source: str,
) -> bool:
    """Return True when a visual asset looks like a chart/directory worth structured extraction."""
    seed = " ".join(
        part for part in (page_text, ocr_text, vision_caption, source) if str(part or "").strip()
    ).lower()
    markers = {
        "organization chart",
        "organisational chart",
        "org chart",
        "directorate",
        "department",
        "ministry",
        "staff directory",
        "directory",
        "hierarchy",
        "reporting line",
        "office structure",
        "general directorate",
        "org structure",
    }
    return any(marker in seed for marker in markers)


def _extract_structured_visual_notes(
    image_bytes: bytes,
    *,
    source: str,
    page_number: str,
    vision_caption: str,
    ocr_text: str,
) -> str:
    """Extract retrieval-friendly structured notes for org charts/directories when possible."""
    raw_model = os.getenv("RAG_VISION_MODEL", "").strip()
    if not raw_model:
        return ""

    provider, model = resolve_provider_and_model(raw_model)
    prompt = (
        "Analyze this image for retrieval. "
        "If it is an organizational chart, staff directory, roster, hierarchy, or office structure, "
        "extract only what is visibly legible in compact prose: document type, top heading, top-most role/title, "
        "top-most person name if readable, other visible names/titles, and any obvious reporting hierarchy. "
        "Preserve original spellings when legible. Mention the visible script or language if recognizable. "
        "If names or titles are not readable, say that clearly. "
        "Do not speculate.\n\n"
        f"Source: {source}\n"
        f"Page: {page_number}\n"
        f"Existing vision caption: {vision_caption or 'none'}\n"
        f"Existing OCR text: {ocr_text or 'none'}"
    )

    try:
        if provider == "llama_cpp":
            chat_format = os.getenv("RAG_VISION_CHAT_FORMAT", "").strip()
            llama_cpp_n_ctx = int(os.getenv("RAG_VISION_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
            llama_cpp_n_batch = int(os.getenv("RAG_VISION_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
            llama_cpp_n_gpu_layers = int(os.getenv("RAG_VISION_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
            llama_cpp_flash_attn = os.getenv(
                "RAG_VISION_FLASH_ATTN",
                os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
            ).strip().lower() in {"1", "true", "yes", "on"}
            vision_model = _get_rag_vision_model(
                raw_model,
                "",
                llama_cpp_n_ctx,
                llama_cpp_n_batch,
                llama_cpp_n_gpu_layers,
                llama_cpp_flash_attn,
                chat_format,
            )
            return vision_model.caption(image_bytes, prompt)[:1600]

        if provider == "llama_server":
            base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
            try:
                from langchain_openai import ChatOpenAI
            except Exception:
                return ""
            llm = ChatOpenAI(
                base_url=base_url,
                api_key="llama-server",
                model=model or os.getenv("LLAMA_SERVER_MODEL", "local").strip() or "local",
                temperature=0.2,
            )
            data_uri = "data:image/png;base64," + base64.b64encode(image_bytes).decode("ascii")
            response = llm.invoke([
                {
                    "role": "system",
                    "content": "You analyze images for retrieval and return compact factual notes.",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_uri}},
                    ],
                },
            ])
            return _extract_text_from_llm_response(response)[:1600]

        import httpx

        base_url = os.getenv("RAG_OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "images": [base64.b64encode(image_bytes).decode("ascii")],
        }
        with httpx.Client(timeout=45.0) as client:
            response = client.post(f"{base_url}/api/generate", json=payload)
            response.raise_for_status()
            data = response.json()
    except Exception:
        return ""

    notes = re.sub(r"\s+", " ", str(data.get("response", "") or "").strip())
    return notes[:1600]


def _assess_image_asset_quality(image_bytes: bytes) -> dict[str, Any]:
    """Estimate whether an extracted image is substantial enough for RAG indexing."""
    result: dict[str, Any] = {
        "keep": False,
        "reason": "unreadable",
        "width": 0,
        "height": 0,
        "pixel_count": 0,
        "aspect": "",
    }
    try:
        from PIL import Image

        with Image.open(BytesIO(image_bytes)) as image:
            width, height = image.size
            pixel_count = width * height
            aspect_ratio = (max(width, height) / max(1, min(width, height))) if min(width, height) else 0.0
            result.update(
                {
                    "width": width,
                    "height": height,
                    "pixel_count": pixel_count,
                    "aspect": "wide" if width > height * 1.2 else "tall" if height > width * 1.2 else "square-ish",
                }
            )

            min_edge = int(os.getenv("RAG_MIN_IMAGE_EDGE", "80"))
            min_pixels = int(os.getenv("RAG_MIN_IMAGE_PIXELS", "12000"))
            max_banner_aspect = float(os.getenv("RAG_MAX_DECORATIVE_ASPECT", "12"))

            if width < min_edge or height < min_edge:
                result["reason"] = "too-small"
                return result
            if pixel_count < min_pixels:
                result["reason"] = "too-few-pixels"
                return result
            if aspect_ratio >= max_banner_aspect:
                result["reason"] = "decorative-aspect"
                return result

            result["keep"] = True
            result["reason"] = "ok"
            return result
    except Exception:
        return result


def _build_visual_summary(
    image_bytes: bytes,
    *,
    source: str,
    page_number: str,
    image_index: int,
    page_text: str,
    ocr_text: str,
    vision_caption: str,
    structured_notes: str = "",
) -> tuple[str, str]:
    """Build retrieval text and topics for an extracted visual asset."""
    cleaned_page_text = re.sub(r"\s+", " ", (page_text or "").strip())
    page_snippet = cleaned_page_text[:500]
    ocr_snippet = re.sub(r"\s+", " ", (ocr_text or "").strip())[:500]
    vision_snippet = re.sub(r"\s+", " ", (vision_caption or "").strip())[:500]
    structured_snippet = re.sub(r"\s+", " ", (structured_notes or "").strip())[:700]
    topics_seed = " ".join(
        part for part in (structured_snippet, vision_snippet, ocr_snippet, cleaned_page_text[:1500]) if part
    )
    topics = _extract_key_topics(topics_seed) if topics_seed else []
    topics_str = ", ".join(topics)

    summary_parts = [
        f"Extracted image {image_index} from {source}",
        f"page {page_number}" if page_number else "page unknown",
    ]

    try:
        from PIL import Image

        with Image.open(BytesIO(image_bytes)) as image:
            width, height = image.size
            aspect = "wide" if width > height * 1.2 else "tall" if height > width * 1.2 else "square-ish"
            summary_parts.append(
                f"visual characteristics: {width}x{height}px, {aspect}, mode {image.mode}"
            )
    except Exception:
        pass

    if topics_str:
        summary_parts.append(f"topics: {topics_str}")
    if vision_snippet:
        summary_parts.append(f"vision caption: {vision_snippet}")
    if structured_snippet:
        summary_parts.append(f"structured visual notes: {structured_snippet}")
    if ocr_snippet:
        summary_parts.append(f"ocr text: {ocr_snippet}")
    if page_snippet:
        summary_parts.append(f"nearby page text: {page_snippet}")

    retrieval_text = ". ".join(summary_parts)
    return retrieval_text, topics_str


def _extract_video_frame_records(
    file_path: Path,
    *,
    project_name: str,
    theme_name: str,
    source_fingerprint: str,
    date_added: str,
    doc_meta: dict[str, str],
    extra_meta: dict[str, str] | None = None,
) -> tuple[list[str], list[dict[str, str]], list[str]]:
    """Extract sampled video frames and convert them into image retrieval records."""
    docs_payload: list[str] = []
    metas_payload: list[dict[str, str]] = []
    ids_payload: list[str] = []

    enabled = os.getenv("RAG_ENABLE_VIDEO_FRAMES", "true").strip().lower()
    if enabled not in {"1", "true", "yes", "on"}:
        return docs_payload, metas_payload, ids_payload

    ffmpeg_exe = _get_ffmpeg_executable()
    asset_dir = RAG_ASSET_DIR / source_fingerprint
    frame_dir = asset_dir / "video_frames"
    frame_dir.mkdir(parents=True, exist_ok=True)
    interval_seconds = max(1, int(os.getenv("RAG_VIDEO_FRAME_INTERVAL_SECONDS", "10")))
    max_frames = max(1, int(os.getenv("RAG_VIDEO_MAX_FRAMES", "12")))

    output_pattern = str(frame_dir / "frame_%04d.jpg")
    command = [
        ffmpeg_exe,
        "-y",
        "-i",
        str(file_path),
        "-vf",
        f"fps=1/{interval_seconds}",
        "-frames:v",
        str(max_frames),
        output_pattern,
    ]
    result = subprocess.run(command, capture_output=True, text=True)
    if result.returncode != 0:
        return docs_payload, metas_payload, ids_payload

    frame_paths = sorted(frame_dir.glob("frame_*.jpg"))
    for frame_index, frame_path in enumerate(frame_paths, start=1):
        try:
            image_bytes = frame_path.read_bytes()
        except Exception:
            continue

        quality = _assess_image_asset_quality(image_bytes)
        if not quality.get("keep"):
            continue

        timecode_seconds = (frame_index - 1) * interval_seconds
        timecode_label = f"{timecode_seconds}s"
        vision_caption = _extract_image_vision_caption(image_bytes)
        ocr_text = _extract_image_ocr_text(image_bytes)
        structured_notes = ""
        if _should_extract_structured_visual_notes("", ocr_text, vision_caption, file_path.name):
            structured_notes = _extract_structured_visual_notes(
                image_bytes,
                source=file_path.name,
                page_number=timecode_label,
                vision_caption=vision_caption,
                ocr_text=ocr_text,
            )
        retrieval_text, topics_str = _build_visual_summary(
            image_bytes,
            source=file_path.name,
            page_number=timecode_label,
            image_index=frame_index,
            page_text="",
            ocr_text=ocr_text,
            vision_caption=vision_caption,
            structured_notes=structured_notes,
        )
        section = f"Video frame {timecode_label}"
        chunk_fingerprint = _build_chunk_fingerprint(
            retrieval_text,
            index=frame_index,
            section=section,
            page_number=timecode_label,
        )
        chunk_meta: dict[str, str] = {
            "project": project_name,
            "source": file_path.name,
            "theme": theme_name,
            "chunking_method": "multimodal_video_frame",
            "file_path": str(file_path),
            "source_fingerprint": source_fingerprint,
            "chunk_fingerprint": chunk_fingerprint,
            "section": section,
            "topics": topics_str[:200],
            "date_added": date_added,
            "page_number": timecode_label,
            "modality": "image",
            "asset_path": _safe_asset_relative_path(frame_path),
            "image_filter_reason": str(quality.get("reason", "")),
            "image_width": str(quality.get("width", "")),
            "image_height": str(quality.get("height", "")),
            "image_pixels": str(quality.get("pixel_count", "")),
            "vision_caption": vision_caption[:500],
            "vision_caption_source": "llama_cpp" if vision_caption else "",
            "ocr_text": ocr_text[:500],
            "structured_visual_notes": structured_notes[:700],
            "visual_summary": retrieval_text[:500],
            "doc_title": doc_meta.get("doc_title", "")[:300],
            "doc_authors": doc_meta.get("doc_authors", "")[:300],
            "doc_year": doc_meta.get("doc_year", ""),
        }
        if extra_meta:
            for k, v in extra_meta.items():
                chunk_meta[k] = str(v)[:300]

        docs_payload.append(retrieval_text)
        metas_payload.append(chunk_meta)
        ids_payload.append(
            _build_chunk_id(
                project=project_name,
                theme=theme_name,
                source_fingerprint=source_fingerprint,
                chunk_fingerprint=chunk_fingerprint,
            )
        )

    return docs_payload, metas_payload, ids_payload


def _extract_standalone_image_records(
    file_path: Path,
    *,
    project_name: str,
    theme_name: str,
    source_fingerprint: str,
    date_added: str,
    doc_meta: dict[str, str],
    extra_meta: dict[str, str] | None = None,
) -> tuple[list[str], list[dict[str, str]], list[str]]:
    """Convert a standalone image file into a multimodal image retrieval record."""
    docs_payload: list[str] = []
    metas_payload: list[dict[str, str]] = []
    ids_payload: list[str] = []

    try:
        image_bytes = file_path.read_bytes()
    except Exception:
        return docs_payload, metas_payload, ids_payload

    quality = _assess_image_asset_quality(image_bytes)
    if not quality.get("keep"):
        return docs_payload, metas_payload, ids_payload

    asset_dir = RAG_ASSET_DIR / source_fingerprint
    asset_dir.mkdir(parents=True, exist_ok=True)
    asset_path = asset_dir / file_path.name
    if not asset_path.exists():
        asset_path.write_bytes(image_bytes)

    vision_caption = _extract_image_vision_caption(image_bytes)
    ocr_text = _extract_image_ocr_text(image_bytes)
    structured_notes = ""
    if _should_extract_structured_visual_notes("", ocr_text, vision_caption, file_path.name):
        structured_notes = _extract_structured_visual_notes(
            image_bytes,
            source=file_path.name,
            page_number="1",
            vision_caption=vision_caption,
            ocr_text=ocr_text,
        )
    retrieval_text, topics_str = _build_visual_summary(
        image_bytes,
        source=file_path.name,
        page_number="1",
        image_index=1,
        page_text="",
        ocr_text=ocr_text,
        vision_caption=vision_caption,
        structured_notes=structured_notes,
    )
    section = "Standalone image"
    chunk_fingerprint = _build_chunk_fingerprint(
        retrieval_text,
        index=1,
        section=section,
        page_number="1",
    )
    chunk_meta: dict[str, str] = {
        "project": project_name,
        "source": file_path.name,
        "theme": theme_name,
        "chunking_method": "multimodal_image_file",
        "file_path": str(file_path),
        "source_fingerprint": source_fingerprint,
        "chunk_fingerprint": chunk_fingerprint,
        "section": section,
        "topics": topics_str[:200],
        "date_added": date_added,
        "page_number": "1",
        "modality": "image",
        "asset_path": _safe_asset_relative_path(asset_path),
        "image_filter_reason": str(quality.get("reason", "")),
        "image_width": str(quality.get("width", "")),
        "image_height": str(quality.get("height", "")),
        "image_pixels": str(quality.get("pixel_count", "")),
        "vision_caption": vision_caption[:500],
        "vision_caption_source": "llama_cpp" if vision_caption else "",
        "ocr_text": ocr_text[:500],
        "structured_visual_notes": structured_notes[:700],
        "visual_summary": retrieval_text[:500],
        "doc_title": doc_meta.get("doc_title", "")[:300],
        "doc_authors": doc_meta.get("doc_authors", "")[:300],
        "doc_year": doc_meta.get("doc_year", ""),
    }
    if extra_meta:
        for k, v in extra_meta.items():
            chunk_meta[k] = str(v)[:300]

    docs_payload.append(retrieval_text)
    metas_payload.append(chunk_meta)
    ids_payload.append(
        _build_chunk_id(
            project=project_name,
            theme=theme_name,
            source_fingerprint=source_fingerprint,
            chunk_fingerprint=chunk_fingerprint,
        )
    )
    return docs_payload, metas_payload, ids_payload


def _extract_table_candidates(page_text: str) -> list[str]:
    """Extract table-like text blocks from a page using simple layout heuristics."""
    raw_lines = [(line or "").strip() for line in (page_text or "").splitlines()]
    table_blocks: list[list[str]] = []
    current_block: list[tuple[str, list[str]]] = []

    def _split_tabular_columns(line: str) -> list[str]:
        if not line:
            return []
        if "|" in line:
            cells = [cell.strip() for cell in line.split("|")]
        elif "\t" in line:
            cells = [cell.strip() for cell in line.split("\t")]
        else:
            cells = [cell.strip() for cell in re.split(r"\s{2,}", line)]
        return [cell for cell in cells if cell]

    def _line_looks_tabular(line: str) -> tuple[bool, list[str]]:
        cells = _split_tabular_columns(line)
        if len(cells) < 2:
            return False, []

        has_explicit_delimiter = "|" in line or "\t" in line
        has_column_gaps = bool(re.search(r"\S\s{2,}\S", line))
        short_cell_count = sum(1 for cell in cells if len(cell) <= 40)
        mostly_sentence_row = len(cells) == 2 and all(len(cell.split()) >= 5 for cell in cells)

        looks_tabular = (
            has_explicit_delimiter
            or (has_column_gaps and short_cell_count >= 2 and not mostly_sentence_row)
        )
        return looks_tabular, cells

    def _flush_current_block() -> None:
        if len(current_block) < 2:
            return
        column_counts = [len(cells) for _line, cells in current_block]
        max_columns = max(column_counts, default=0)
        min_columns = min(column_counts, default=0)
        if min_columns < 2:
            return
        if max_columns >= 3 or max_columns == min_columns:
            table_blocks.append([line for line, _cells in current_block])

    for line in raw_lines:
        looks_tabular, cells = _line_looks_tabular(line)
        if looks_tabular:
            current_block.append((line, cells))
        else:
            _flush_current_block()
            current_block = []

    _flush_current_block()

    normalized_blocks: list[str] = []
    for block in table_blocks:
        normalized = "\n".join(block).strip()
        if normalized and len(normalized) >= 20:
            normalized_blocks.append(normalized[:2500])
    return normalized_blocks


def _build_table_summary(
    table_text: str,
    *,
    source: str,
    page_number: str,
    table_index: int,
    page_text: str,
) -> tuple[str, str]:
    """Build retrieval text and topics for an extracted table candidate."""
    normalized_table = re.sub(r"\s+", " ", (table_text or "").strip())
    cleaned_page_text = re.sub(r"\s+", " ", (page_text or "").strip())
    topics_seed = " ".join(part for part in (normalized_table[:1200], cleaned_page_text[:800]) if part)
    topics = _extract_key_topics(topics_seed) if topics_seed else []
    topics_str = ", ".join(topics)

    summary_parts = [
        f"Extracted table {table_index} from {source}",
        f"page {page_number}" if page_number else "page unknown",
    ]
    if topics_str:
        summary_parts.append(f"topics: {topics_str}")
    if normalized_table:
        summary_parts.append(f"table text: {normalized_table[:800]}")
    if cleaned_page_text:
        summary_parts.append(f"nearby page text: {cleaned_page_text[:400]}")

    retrieval_text = ". ".join(summary_parts)
    return retrieval_text, topics_str


def _normalize_table_rows(rows: list[list[Any]] | None) -> str:
    """Convert extracted table rows into a stable tab-separated text block."""
    if not rows:
        return ""

    normalized_lines: list[str] = []
    for row in rows:
        if not row:
            continue
        cells = [re.sub(r"\s+", " ", str(cell or "").strip()) for cell in row]
        if any(cells):
            normalized_lines.append("\t".join(cells))
    return "\n".join(normalized_lines).strip()


def _extract_tables_from_pdf_page(pdf_page: Any) -> list[str]:
    """Extract parser-backed table text blocks from a PyMuPDF page."""
    try:
        finder = pdf_page.find_tables()
    except Exception:
        return []

    tables = getattr(finder, "tables", None) or []
    extracted_tables: list[str] = []
    for table in tables:
        rows: list[list[Any]] | None = None
        try:
            rows = table.extract()
        except Exception:
            rows = None

        normalized = _normalize_table_rows(rows)
        if normalized and len(normalized) >= 20:
            extracted_tables.append(normalized[:2500])
    return extracted_tables


def _extract_pdf_table_records(
    file_path: Path,
    pages: list[Any],
    *,
    project_name: str,
    theme_name: str,
    source_fingerprint: str,
    date_added: str,
    doc_meta: dict[str, str],
    extra_meta: dict[str, str] | None = None,
) -> tuple[list[str], list[dict[str, str]], list[str]]:
    """Extract table-like text regions from PDF pages and convert them into retrieval records."""
    docs_payload: list[str] = []
    metas_payload: list[dict[str, str]] = []
    ids_payload: list[str] = []

    asset_dir = RAG_ASSET_DIR / source_fingerprint
    asset_dir.mkdir(parents=True, exist_ok=True)

    pdf_pages_by_number: dict[int, Any] = {}
    try:
        import fitz

        pdf = fitz.open(str(file_path))
    except Exception:
        pdf = None

    if pdf is not None:
        try:
            for page_index in range(len(pdf)):
                try:
                    pdf_pages_by_number[page_index + 1] = pdf.load_page(page_index)
                except Exception:
                    continue
        finally:
            try:
                pdf.close()
            except Exception:
                pass

    for page_idx, page_doc in enumerate(pages, start=1):
        page_number = str(page_idx)
        page_text = getattr(page_doc, "page_content", "") or ""
        extraction_method = "heuristic"
        table_blocks = _extract_tables_from_pdf_page(pdf_pages_by_number.get(page_idx))
        if table_blocks:
            extraction_method = "pymupdf"
        else:
            table_blocks = _extract_table_candidates(page_text)

        for table_index, table_text in enumerate(table_blocks, start=1):
            retrieval_text, topics_str = _build_table_summary(
                table_text,
                source=file_path.name,
                page_number=page_number,
                table_index=table_index,
                page_text=page_text,
            )
            section = f"Table asset page {page_number}"
            chunk_fingerprint = _build_chunk_fingerprint(
                retrieval_text,
                index=page_idx * 1000 + table_index,
                section=section,
                page_number=page_number,
            )
            asset_path = asset_dir / f"page_{page_number}_table_{table_index}.txt"
            if not asset_path.exists():
                asset_path.write_text(table_text, encoding="utf-8")

            chunk_meta: dict[str, str] = {
                "project": project_name,
                "source": file_path.name,
                "theme": theme_name,
                "chunking_method": "multimodal_table",
                "file_path": str(file_path),
                "source_fingerprint": source_fingerprint,
                "chunk_fingerprint": chunk_fingerprint,
                "section": section,
                "table_extraction_method": extraction_method,
                "topics": topics_str[:200],
                "date_added": date_added,
                "page_number": page_number,
                "modality": "table",
                "asset_path": _safe_asset_relative_path(asset_path),
                "table_text": table_text[:1000],
                "table_summary": retrieval_text[:500],
                "doc_title": doc_meta.get("doc_title", "")[:300],
                "doc_authors": doc_meta.get("doc_authors", "")[:300],
                "doc_year": doc_meta.get("doc_year", ""),
            }
            if extra_meta:
                for k, v in extra_meta.items():
                    chunk_meta[k] = str(v)[:300]

            docs_payload.append(retrieval_text)
            metas_payload.append(chunk_meta)
            ids_payload.append(
                _build_chunk_id(
                    project=project_name,
                    theme=theme_name,
                    source_fingerprint=source_fingerprint,
                    chunk_fingerprint=chunk_fingerprint,
                )
            )

    return docs_payload, metas_payload, ids_payload


def _extract_pdf_image_records(
    file_path: Path,
    pages: list[Any],
    *,
    project_name: str,
    theme_name: str,
    source_fingerprint: str,
    date_added: str,
    doc_meta: dict[str, str],
    extra_meta: dict[str, str] | None = None,
) -> tuple[list[str], list[dict[str, str]], list[str]]:
    """Extract embedded PDF images and convert them into retrieval records."""
    docs_payload: list[str] = []
    metas_payload: list[dict[str, str]] = []
    ids_payload: list[str] = []

    try:
        import fitz
    except Exception:
        return docs_payload, metas_payload, ids_payload

    asset_dir = RAG_ASSET_DIR / source_fingerprint
    asset_dir.mkdir(parents=True, exist_ok=True)

    try:
        pdf = fitz.open(str(file_path))
    except Exception:
        return docs_payload, metas_payload, ids_payload

    try:
        for page_idx in range(len(pdf)):
            page = pdf.load_page(page_idx)
            page_number = str(page_idx + 1)
            page_text = ""
            if page_idx < len(pages):
                page_text = getattr(pages[page_idx], "page_content", "") or ""
            image_list = page.get_images(full=True)
            for image_index, image_info in enumerate(image_list, start=1):
                xref = image_info[0]
                try:
                    extracted = pdf.extract_image(xref)
                except Exception:
                    continue
                image_bytes = extracted.get("image")
                image_ext = extracted.get("ext", "png")
                if not image_bytes:
                    continue
                quality = _assess_image_asset_quality(image_bytes)
                if not quality.get("keep"):
                    continue

                asset_path = asset_dir / f"page_{page_number}_image_{image_index}.{image_ext}"
                if not asset_path.exists():
                    asset_path.write_bytes(image_bytes)

                vision_caption = _extract_image_vision_caption(image_bytes)
                ocr_text = _extract_image_ocr_text(image_bytes)
                structured_notes = ""
                if _should_extract_structured_visual_notes(page_text, ocr_text, vision_caption, file_path.name):
                    structured_notes = _extract_structured_visual_notes(
                        image_bytes,
                        source=file_path.name,
                        page_number=page_number,
                        vision_caption=vision_caption,
                        ocr_text=ocr_text,
                    )
                retrieval_text, topics_str = _build_visual_summary(
                    image_bytes,
                    source=file_path.name,
                    page_number=page_number,
                    image_index=image_index,
                    page_text=page_text,
                    ocr_text=ocr_text,
                    vision_caption=vision_caption,
                    structured_notes=structured_notes,
                )
                section = f"Visual asset page {page_number}"
                chunk_fingerprint = _build_chunk_fingerprint(
                    retrieval_text,
                    index=page_idx * 1000 + image_index,
                    section=section,
                    page_number=page_number,
                )
                chunk_meta: dict[str, str] = {
                    "project": project_name,
                    "source": file_path.name,
                    "theme": theme_name,
                    "chunking_method": "multimodal_image",
                    "file_path": str(file_path),
                    "source_fingerprint": source_fingerprint,
                    "chunk_fingerprint": chunk_fingerprint,
                    "section": section,
                    "topics": topics_str[:200],
                    "date_added": date_added,
                    "page_number": page_number,
                    "modality": "image",
                    "asset_path": _safe_asset_relative_path(asset_path),
                    "image_filter_reason": str(quality.get("reason", "")),
                    "image_width": str(quality.get("width", "")),
                    "image_height": str(quality.get("height", "")),
                    "image_pixels": str(quality.get("pixel_count", "")),
                    "vision_caption": vision_caption[:500],
                    "vision_caption_source": "llama_cpp" if vision_caption else "",
                    "ocr_text": ocr_text[:500],
                    "structured_visual_notes": structured_notes[:700],
                    "visual_summary": retrieval_text[:500],
                    "doc_title": doc_meta.get("doc_title", "")[:300],
                    "doc_authors": doc_meta.get("doc_authors", "")[:300],
                    "doc_year": doc_meta.get("doc_year", ""),
                }
                if extra_meta:
                    for k, v in extra_meta.items():
                        chunk_meta[k] = str(v)[:300]

                docs_payload.append(retrieval_text)
                metas_payload.append(chunk_meta)
                ids_payload.append(
                    _build_chunk_id(
                        project=project_name,
                        theme=theme_name,
                        source_fingerprint=source_fingerprint,
                        chunk_fingerprint=chunk_fingerprint,
                    )
                )
    finally:
        pdf.close()

    return docs_payload, metas_payload, ids_payload


def _infer_chunking_method_from_document(
    document_text: str,
    metadata: dict[str, Any] | None = None,
) -> str:
    """Best-effort recovery of chunking method for older chunks without metadata."""
    meta = metadata or {}
    method = (meta.get("chunking_method") or "").strip()
    if method:
        return method

    text = (document_text or "").strip()
    has_context = "[Context:" in text
    has_late = "[LateCtx:" in text
    has_topics = "[Topics:" in text

    if has_context and has_late:
        return "semantic_contextual_late"
    if has_context:
        return "contextual"
    if has_late:
        return "late_chunking"
    if has_topics:
        return "semantic"
    return "recursive"


def _derive_apa_reference(source: str) -> tuple[str, str]:
    """Legacy: build a best-effort APA-style reference from a source filename.
    Used only as fallback when harvested document metadata is unavailable.
    """
    stem = Path(source or "unknown").stem
    normalized = re.sub(r"[_\-]+", " ", stem)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    normalized = re.sub(r"^\d+(?:\.\d+)*\s*", "", normalized)

    year_match = re.search(r"\b(19|20)\d{2}\b", normalized)
    year = year_match.group(0) if year_match else "n.d."

    title = normalized if normalized else "Untitled document"
    apa_ref = f"{title}. ({year}). [PDF]."
    return apa_ref, year


def _derive_intext_author(source: str) -> str:
    """Legacy: build a short in-text author token from source filename.
    Returns a clearly-labelled filename token so citations are never
    mistaken for real author names when no metadata is available.
    """
    stem = Path(source or "unknown").stem
    normalized = re.sub(r"[_\-]+", " ", stem)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    normalized = re.sub(r"^\d+(?:\.\d+)*\s*", "", normalized)
    if not normalized:
        return "[Author unknown]"
    # Return the full normalized stem (capped) so it is recognisable as a
    # filename, not fabricated as a real author surname.
    return normalized[:60]


def _harvest_pdf_metadata(file_path: Path) -> dict[str, str]:
    """Extract real document metadata from a PDF using PyMuPDF.

    Returns a dict with keys: doc_title, doc_authors, doc_year.
    Falls back to empty strings when metadata is absent or unreadable.
    """
    meta: dict[str, str] = {"doc_title": "", "doc_authors": "", "doc_year": ""}
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(file_path))
        pdf_meta = doc.metadata or {}
        doc.close()

        title = (pdf_meta.get("title") or "").strip()
        author = (pdf_meta.get("author") or "").strip()
        creation_date = (pdf_meta.get("creationDate") or "").strip()

        # creationDate format: D:YYYYMMDDHHmmSS or D:YYYY...
        year_match = re.search(r"(19|20)\d{2}", creation_date)
        year = year_match.group(0) if year_match else ""

        # Also try to find year in title/filename if not in date
        if not year:
            for candidate in [title, file_path.stem]:
                y = re.search(r"\b(19|20)\d{2}\b", candidate)
                if y:
                    year = y.group(0)
                    break

        meta["doc_title"] = title
        meta["doc_authors"] = author
        meta["doc_year"] = year
    except Exception:
        pass
    return meta


def _build_apa_reference(chunk_meta: dict[str, Any]) -> tuple[str, str]:
    """Build an APA reference from chunk metadata.

    Uses harvested PDF metadata (doc_title, doc_authors, doc_year, page_number)
    when available, falling back to filename-derived heuristics.

    Returns (apa_ref_string, year).
    """
    doc_title = chunk_meta.get("doc_title", "").strip()
    doc_authors = chunk_meta.get("doc_authors", "").strip()
    doc_year = chunk_meta.get("doc_year", "").strip()
    source = chunk_meta.get("source", "unknown")

    if doc_title or doc_authors:
        # Use real metadata
        year = doc_year or "n.d."
        author_part = doc_authors if doc_authors else "[Author unknown]"
        title_part = doc_title if doc_title else source
        apa_ref = f"{author_part}. ({year}). {title_part}. [PDF]."
        return apa_ref, year

    # Fallback to filename heuristic
    return _derive_apa_reference(source)


def _build_intext_author(chunk_meta: dict[str, Any]) -> str:
    """Build a short in-text author token from chunk metadata.

    Uses harvested PDF author metadata, falling back to filename heuristic.
    """
    doc_authors = chunk_meta.get("doc_authors", "").strip()
    if doc_authors:
        # Use first surname from author string (e.g. "Smith, J.; Jones, A." → "Smith")
        first_author = re.split(r"[;,]", doc_authors)[0].strip()
        return first_author[:40] if first_author else "Unknown"
    return _derive_intext_author(chunk_meta.get("source", "unknown"))


def get_vector_collection(project: str = "Default"):
    """Get the RAG vector collection for a project, used by tools and UI."""
    return _get_vector_collection(project)


def _extract_section_headings(text: str) -> list[tuple[int, str]]:
    """Extract section/subsection headings with their character positions.

    Detects:
    - Markdown headings (# Heading)
    - Uppercase lines followed by newlines (SECTION TITLE)
    - Lines ending with colon that look like headings (Section Title:)
    - Numbered sections (1. Section, 1.1 Subsection)

    Falls back to positional heuristics when no structural headings are found
    (e.g., scanned PDFs, unstructured documents):
    - Short lines (≤80 chars) that are mostly capitalised (ALL_CAPS ratio > 0.4)
    - First 100 chars of the full text treated as a likely title
    """
    headings: list[tuple[int, str]] = []
    for m in re.finditer(
        r'^(?:'
        r'#{1,4}\s+(.+)'           # Markdown headings
        r'|(\d+(?:\.\d+)*)\s+(.+)'  # Numbered sections like "1.1 Introduction"
        r'|([A-Z][A-Z\s]{4,})\s*$'  # ALL CAPS lines (min 5 chars)
        r'|(.{5,80}):\s*$'          # Lines ending with colon
        r')$',
        text,
        re.MULTILINE,
    ):
        heading = (
            m.group(1)
            or (f"{m.group(2)} {m.group(3)}" if m.group(2) else None)
            or m.group(4)
            or m.group(5)
        )
        if heading:
            heading = heading.strip()
            if len(heading) > 3 and not heading.endswith('.'):
                headings.append((m.start(), heading))

    # ── Positional fallback for documents with no detectable structure ──────
    if not headings:
        # Treat any short line (≤80 chars) with high uppercase ratio as a heading
        for m in re.finditer(r'^(.{5,80})$', text, re.MULTILINE):
            line = m.group(1).strip()
            if not line:
                continue
            alpha_chars = [c for c in line if c.isalpha()]
            if alpha_chars and sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars) >= 0.4:
                headings.append((m.start(), line))

        # If still nothing, use first 100 chars as a document-level title
        if not headings:
            title_candidate = text[:100].strip().replace("\n", " ")
            if title_candidate:
                headings.append((0, title_candidate[:80]))

    return headings


def _get_heading_context(position: int, headings: list[tuple[int, str]], max_depth: int = 3) -> str:
    """Get the heading hierarchy active at a given character position.

    Returns a string like "Chapter 1 > Introduction > Background"
    """
    active = []
    for hpos, htitle in headings:
        if hpos <= position:
            # Replace deeper heading if it appears after a shallower one
            active.append(htitle)
        else:
            break
    # Keep last max_depth headings as context
    return " > ".join(active[-max_depth:]) if active else ""


def _extract_key_topics(text: str, max_topics: int = 5) -> list[str]:
    """Extract key noun phrases / terms from a chunk.

    Uses KeyBERT (ML-based keyphrase extraction) when available (B3);
    falls back to capitalized-phrase / acronym / quoted-term regex heuristics.
    """
    try:
        kw_model = _get_keybert_model()
        keywords = kw_model.extract_keywords(
            text[:3000],
            keyphrase_ngram_range=(1, 2),
            stop_words="english",
            top_n=max_topics,
        )
        topics = [kw.strip() for kw, _ in keywords if kw.strip()]
        if topics:
            return topics
    except Exception:
        pass

    # Fallback: regex heuristic
    clean = re.sub(r'\s+', ' ', text)
    phrases = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b', clean)
    acronyms = re.findall(r'\b([A-Z]{2,6})\b', clean)
    quoted = re.findall(r'"([^"]{3,40})"', clean)
    seen: set[str] = set()
    topics_fallback: list[str] = []
    for t in phrases + acronyms + quoted:
        t_clean = t.strip()
        t_lower = t_clean.lower()
        if t_lower not in seen and len(t_lower) > 2:
            seen.add(t_lower)
            topics_fallback.append(t_clean)
    return topics_fallback[:max_topics]


# Module-level cache for context LLM — invalidates when model/provider changes
_context_llm_cache: dict[str, Any] = {}  # {"cache_key": str, "model": str, "llm": Any}

# Last error/status from _generate_chunk_context — use a dict so Streamlit imports
# stay live (string re-binding doesn't update an already-imported name).
_context_llm_status: dict[str, str] = {
    "error": "",
    "model": "",
    "state": "idle",
}
_last_rag_query_diagnostics: dict[str, Any] = {}


def _strip_ollama_provider_prefix(model: str) -> str:
    """Strip LangChain provider prefix (e.g. 'ollama:') that Ollama's API rejects."""
    if ":" in model:
        provider, _, rest = model.partition(":")
        if provider.lower() in ("ollama", "openai", "anthropic", "google"):
            return rest
    return model


def _get_context_llm():
    """Return a cached context LLM for generating chunk context (contextual chunking).

    Re-creates the client when RAG_CONTEXT_LLM_MODEL changes at runtime
    (unlike @lru_cache which permanently caches the first value).
    """
    raw = os.getenv("RAG_CONTEXT_LLM_MODEL", os.getenv("DEEPAGENT_MODEL", "gemma4:26b"))
    provider, model = resolve_provider_and_model(raw)
    cache_key = f"{provider}:{model}"
    if _context_llm_cache.get("cache_key") != cache_key:
        _context_llm_cache["cache_key"] = cache_key
        _context_llm_cache["model"] = model
        _context_llm_cache["llm"] = create_chat_model(model_name=raw, temperature=0)
    _context_llm_status["model"] = cache_key
    return _context_llm_cache["llm"]


@lru_cache(maxsize=1)
def _get_keybert_model():
    """Load and cache a KeyBERT model for ML-based keyphrase extraction (B3)."""
    from keybert import KeyBERT

    return KeyBERT()


def _extract_text_from_llm_response(response: Any) -> str:
    """Robustly extract plain text from a chat-model response.

    Handles:
    - Plain string content (standard models)
    - List content with thinking blocks (Gemma4, DeepSeek-R1, etc.)
      e.g. [{"type": "thinking", "thinking": "..."}, {"type": "text", "text": "..."}]
    - Fallback to str(response)
    """
    return _extract_text_from_content(getattr(response, "content", None) if not isinstance(response, dict) else response.get("content"))


def _extract_text_from_content(content: Any) -> str:
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        parts = [
            block.get("text", "")
            for block in content
            if isinstance(block, dict) and block.get("type") == "text"
        ]
        joined = " ".join(p.strip() for p in parts if p.strip())
        if joined:
            return joined
        for block in content:
            if isinstance(block, dict):
                for v in block.values():
                    if isinstance(v, str) and v.strip():
                        return v.strip()
        return ""
    return str(content).strip() if content is not None else ""


def _generate_chunk_context(full_doc_text: str, chunk_text: str) -> str:
    """Generate a brief situating context for a chunk using Anthropic's contextual chunking prompt.

    Prepends the LLM-generated context to the chunk embedding so retrieval
    is aware of where the chunk sits within the whole document.
    Returns an empty string on any failure (caller falls back to topic tags);
    stores error/status in ``_context_llm_status`` for UI diagnostics.
    """
    prompt = (
        "<document>\n"
        f"{full_doc_text[:8000]}\n"
        "</document>\n"
        "Here is the chunk we want to situate within the whole document:\n"
        "<chunk>\n"
        f"{chunk_text}\n"
        "</chunk>\n"
        "Please give a short succinct context to situate this chunk within the overall document "
        "for the purposes of improving search retrieval of the chunk. "
        "Answer only with the succinct context and nothing else."
    )
    try:
        llm = _get_context_llm()
        response = llm.invoke(prompt)
        ctx = _extract_text_from_llm_response(response)
        if ctx:
            _context_llm_status["error"] = ""
            _context_llm_status["state"] = "context-generated"
        else:
            _context_llm_status["error"] = (
                f"Model '{_context_llm_status['model']}' returned an empty response. "
                "If this is a thinking model, check that Ollama supports its output format."
            )
            _context_llm_status["state"] = "empty-response"
        return ctx[:500]
    except Exception as exc:
        _context_llm_status["error"] = str(exc)
        _context_llm_status["state"] = "llm-error"
        return ""


def _build_splitter(
    chunking_method: str,
    chunk_size: int,
    chunk_overlap: int,
    breakpoint_threshold: float,
):
    """Return the appropriate text splitter based on chosen method.

    Tries SemanticChunker (topic-boundary splitting via Ollama embeddings)
    and falls back silently to RecursiveCharacterTextSplitter if unavailable.

    Methods:
    - ``semantic``: SemanticChunker on topic-similarity breakpoints.
    - ``contextual``: SemanticChunker + LLM context prefix per chunk (Anthropic method).
    - ``recursive``: Fixed-size RecursiveCharacterTextSplitter.
    - ``late_chunking``: SemanticChunker splits + surrounding-sentence context window
      injected per chunk before embedding (late-chunking approximation for Ollama).
    - ``semantic_contextual_late``: All three — semantic splits + LLM context description
      + late-chunking context window. Best quality, slowest.
    """
    _uses_semantic = chunking_method in (
        "semantic", "contextual", "late_chunking", "semantic_contextual_late"
    )
    if _uses_semantic:
        try:
            from langchain_experimental.text_splitter import SemanticChunker
            embed_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
            llama_server_base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
            llama_cpp_n_ctx = int(os.getenv("RAG_LLAMA_CPP_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
            llama_cpp_n_batch = int(os.getenv("RAG_LLAMA_CPP_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
            llama_cpp_n_gpu_layers = int(os.getenv("RAG_LLAMA_CPP_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
            llama_cpp_flash_attn = os.getenv(
                "RAG_LLAMA_CPP_FLASH_ATTN",
                os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
            ).strip().lower() in {"1", "true", "yes", "on"}
            embeddings = _get_rag_embedding_function(
                embed_model,
                llama_server_base_url,
                llama_cpp_n_ctx,
                llama_cpp_n_batch,
                llama_cpp_n_gpu_layers,
                llama_cpp_flash_attn,
            )
            return SemanticChunker(
                embeddings,
                breakpoint_threshold_type="percentile",
                breakpoint_threshold_amount=breakpoint_threshold,
            ), chunking_method  # preserves the requested label
        except Exception:
            pass  # Fall through to recursive

    from langchain_text_splitters import RecursiveCharacterTextSplitter
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "?", "!", " ", ""],
    ), "recursive"


def _build_late_chunk_context(splits: list, idx: int, window: int = 2) -> str:
    """Build a context window string from neighbouring chunks for late chunking.

    Concatenates text from up to ``window`` chunks before and after ``idx``.
    This approximates the document-level context awareness that true late chunking
    achieves via token-level embedding pooling — adapted for Ollama which only
    exposes sentence/chunk-level embeddings.

    Returns a short summary string (≤ 400 chars) to prepend as ``[LateCtx: ...]``.
    """
    texts: list[str] = []
    for off in range(-window, window + 1):
        j = idx + off
        if j < 0 or j >= len(splits) or j == idx:
            continue
        neighbour = (splits[j].page_content or "").strip()
        if neighbour:
            # Take first sentence (up to 120 chars) from each neighbour
            first_sentence = re.split(r"(?<=[.!?])\s+", neighbour)[0][:120]
            texts.append(first_sentence)
    combined = " … ".join(texts)
    return combined[:400]


def _ingest_document_records(
    records: list[dict[str, Any]],
    *,
    project: str = "Default",
    theme: str = "",
    chunk_size: int = 1500,
    chunk_overlap: int = 300,
    chunking_method: str = "semantic",
    breakpoint_threshold: float = 95.0,
) -> dict[str, Any]:
    """Ingest prebuilt text records into the RAG vector store with per-record metadata."""
    from datetime import datetime

    project_name = _normalize_project(project)
    theme_name = theme.strip() or "Unspecified"
    if not records:
        return {
            "loaded_files": 0,
            "total_files": 0,
            "added_chunks": 0,
            "project": project_name,
            "theme": theme_name,
            "failures": ["No document records provided."],
        }

    collection = _get_vector_collection(project_name)
    splitter, method_used = _build_splitter(
        chunking_method=chunking_method,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        breakpoint_threshold=breakpoint_threshold,
    )
    try:
        from langchain_core.documents import Document as LCDocument
    except ImportError:
        from langchain.schema import Document as LCDocument

    added_chunks = 0
    loaded_files = 0
    failures: list[str] = []
    date_added = datetime.now().strftime("%Y-%m-%d")

    for record_index, record in enumerate(records, start=1):
        text = str(record.get("text", "") or "").strip()
        if not text:
            failures.append(f"record {record_index} -> empty text")
            continue

        source_label = str(record.get("source", "") or f"web-result-{record_index}").strip()
        extra_meta = dict(record.get("metadata") or {})
        page_docs = [
            LCDocument(
                page_content=text,
                metadata={"page_number": "1", **extra_meta},
            )
        ]
        splits = splitter.split_documents(page_docs)
        if not splits:
            failures.append(f"{source_label} -> no chunks produced after splitting")
            continue

        full_text = "\n\n".join((getattr(d, "page_content", "") or "") for d in page_docs)
        headings = _extract_section_headings(full_text)
        char_offset = 0
        source_fingerprint = str(record.get("source_fingerprint") or _fingerprint_text(f"web:{source_label}"))

        docs_payload: list[str] = []
        metas_payload: list[dict[str, str]] = []
        ids_payload: list[str] = []
        for idx, split in enumerate(splits):
            chunk_text = (split.page_content or "").strip()
            if not chunk_text:
                continue

            page_number = str(split.metadata.get("page_number", "1") or "1")
            chunk_pos = full_text.find(chunk_text[:80], max(0, char_offset - 100))
            if chunk_pos == -1:
                chunk_pos = char_offset
            char_offset = chunk_pos + len(chunk_text)
            heading_ctx = _get_heading_context(chunk_pos, headings)
            chunk_fingerprint = _build_chunk_fingerprint(
                chunk_text,
                index=idx,
                section=heading_ctx,
                page_number=page_number,
            )
            topics = _extract_key_topics(chunk_text)
            topics_str = ", ".join(topics) if topics else ""

            enriched_parts = []
            if method_used in ("contextual", "semantic_contextual_late"):
                ctx = _generate_chunk_context(full_text, chunk_text)
                if ctx:
                    enriched_parts.append(f"[Context: {ctx}]")
                elif topics_str:
                    enriched_parts.append(f"[Topics: {topics_str}]")
            elif topics_str:
                enriched_parts.append(f"[Topics: {topics_str}]")

            if method_used in ("late_chunking", "semantic_contextual_late"):
                late_ctx = _build_late_chunk_context(splits, idx)
                if late_ctx:
                    enriched_parts.append(f"[LateCtx: {late_ctx}]")

            if heading_ctx:
                enriched_parts.append(f"[Section: {heading_ctx}]")
            enriched_parts.append(chunk_text)
            enriched_text = "\n".join(enriched_parts)

            docs_payload.append(enriched_text)
            chunk_meta: dict[str, str] = {
                "project": project_name,
                "source": source_label[:300],
                "theme": theme_name,
                "chunking_method": method_used,
                "file_path": str(record.get("file_path", source_label))[:500],
                "source_fingerprint": source_fingerprint,
                "chunk_fingerprint": chunk_fingerprint,
                "section": heading_ctx[:200] if heading_ctx else "",
                "topics": topics_str[:200],
                "date_added": date_added,
                "page_number": page_number,
                "modality": str(record.get("modality", "text") or "text"),
                # Verbatim source text (no LLM enrichment) — used by rag_retrieve
                # to display grounded, quotable content instead of the embedding-enriched text.
                "chunk_source_text": chunk_text[:4000],
            }
            for key, value in extra_meta.items():
                chunk_meta[str(key)] = str(value)[:500]
            metas_payload.append(chunk_meta)
            ids_payload.append(
                _build_chunk_id(
                    project=project_name,
                    theme=theme_name,
                    source_fingerprint=source_fingerprint,
                    chunk_fingerprint=chunk_fingerprint,
                )
            )

        if not docs_payload:
            failures.append(f"{source_label} -> no non-empty chunks after enrichment")
            continue

        stale_chunks_deleted = _delete_stale_chunks(
            collection=collection,
            source=source_label[:300],
            theme=theme_name,
            source_fingerprint=source_fingerprint,
        )
        _ = stale_chunks_deleted
        collection.upsert(documents=docs_payload, metadatas=metas_payload, ids=ids_payload)
        added_chunks += len(docs_payload)
        loaded_files += 1

    return {
        "loaded_files": loaded_files,
        "total_files": len(records),
        "added_chunks": added_chunks,
        "project": project_name,
        "theme": theme_name,
        "failures": failures,
        "method_used": method_used,
        "stale_chunks_deleted": 0,
        "stopped": False,
    }


def ingest_rag_paths(
    paths: list[Path],
    project: str = "Default",
    theme: str = "",
    chunk_size: int = 1500,
    chunk_overlap: int = 300,
    chunking_method: str = "semantic",
    breakpoint_threshold: float = 95.0,
    extra_meta: dict[str, str] | None = None,
    on_progress: Any | None = None,
    stop_event: Any | None = None,
) -> dict[str, Any]:
    """Ingest local file paths with context-aware chunking.

    Each chunk is enriched with:
    - Section heading hierarchy (e.g., "Chapter 1 > Methods > Sampling")
    - Key topics extracted from the chunk
    - Source metadata (file, project, theme, date)

    Chunking methods:
    - ``semantic``: Uses ``SemanticChunker`` (Ollama embeddings) to split at
      topic-similarity breakpoints. Produces coherent, idea-aligned chunks.
      ``breakpoint_threshold`` (percentile 0-100) controls sensitivity —
      higher = fewer, larger chunks; lower = more, smaller chunks.
    - ``recursive``: Falls back to ``RecursiveCharacterTextSplitter`` using
      ``chunk_size`` / ``chunk_overlap`` character counts.
    - ``contextual``: Semantic splits + LLM-generated context description per chunk
      (Anthropic's contextual retrieval method). Slow — one LLM call per chunk.
    - ``late_chunking``: Semantic splits + neighbouring-chunk context window (±2 chunks)
      prepended as ``[LateCtx: ...]``. Approximates Jina-style late chunking where each
      chunk's embedding is informed by surrounding document context.
    - ``semantic_contextual_late``: All three combined — semantic splits + LLM context
      + late-chunking window. Best retrieval quality, slowest ingestion.

    Args:
        extra_meta: Optional additional metadata fields to store per chunk
            (e.g. doi, journal, pub_date, citation_count from literature search).
        on_progress: Optional callback ``(step, idx, total, message)`` for progress reporting.
        stop_event: Optional ``threading.Event`` — when set, ingestion stops after the current file.

    The heading context and topics are prepended to the chunk text stored in the
    vector DB so that embedding captures the semantic context even for chunks
    that appear mid-section.
    """
    from datetime import datetime

    def _progress(step: str, idx: int, total: int, msg: str) -> None:
        if on_progress:
            on_progress(step, idx, total, msg)

    def _stopped() -> bool:
        return stop_event is not None and stop_event.is_set()

    project_name = _normalize_project(project)
    theme_name = theme.strip() or "Unspecified"

    if not paths:
        return {
            "loaded_files": 0,
            "total_files": 0,
            "added_chunks": 0,
            "project": project_name,
            "theme": theme_name,
            "failures": ["No file paths provided."],
        }

    collection = _get_vector_collection(project_name)
    splitter, method_used = _build_splitter(
        chunking_method=chunking_method,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        breakpoint_threshold=breakpoint_threshold,
    )

    added_chunks = 0
    loaded_files = 0
    failures: list[str] = []
    stale_chunks_deleted = 0
    total_files = len(paths)

    for file_idx, path in enumerate(paths):
        if _stopped():
            break
        if not path.exists():
            failures.append(f"{path} -> file not found")
            continue

        _progress("ingest", file_idx + 1, total_files, f"Loading: {path.name}…")
        try:
            docs, doc_meta = _load_documents(path)
            if not docs:
                failures.append(f"{path} -> no extractable document content found")
                continue

            # Drop empty pages/sections before splitting to avoid empty upsert payloads.
            docs = [d for d in docs if (getattr(d, "page_content", "") or "").strip()]
            if not docs:
                failures.append(
                    f"{path} -> no extractable text found (file may be image-only/scanned)"
                )
                continue

            splits = splitter.split_documents(docs)
            if not splits:
                failures.append(f"{path} -> no chunks produced after splitting")
                continue

            # Build full text for heading extraction
            full_text = "\n\n".join(
                (getattr(d, "page_content", "") or "") for d in docs
            )
            headings = _extract_section_headings(full_text)
            date_added = datetime.now().strftime("%Y-%m-%d")

            # Track character offset for heading context mapping
            # We rebuild offset from splits since langchain doesn't expose source offsets
            char_offset = 0
            source_fingerprint = _build_source_fingerprint(path)

            docs_payload = []
            metas_payload = []
            ids_payload = []
            for idx, split in enumerate(splits):
                text = (split.page_content or "").strip()
                if not text:
                    continue

                # Inherit page number from source page if available
                page_number = str(split.metadata.get("page_number", "") or "")

                # Find this chunk's approximate position in the full text
                chunk_pos = full_text.find(text[:80], max(0, char_offset - 100))
                if chunk_pos == -1:
                    chunk_pos = char_offset
                char_offset = chunk_pos + len(text)

                # Get heading hierarchy for this chunk
                heading_ctx = _get_heading_context(chunk_pos, headings)
                chunk_fingerprint = _build_chunk_fingerprint(
                    text,
                    index=idx,
                    section=heading_ctx,
                    page_number=page_number,
                )

                # Extract key topics from this chunk
                topics = _extract_key_topics(text)
                topics_str = ", ".join(topics) if topics else ""

                # Build enriched text: prepend context/heading/topics
                # so the embedding captures the semantic position of the chunk.
                #
                # contextual              — LLM generates a description of the chunk
                #                          in context of the full document (Anthropic method).
                # late_chunking           — surrounding chunks' first sentences prepended
                #                          as [LateCtx: ...] to simulate token-pooling
                #                          context awareness (Jina-style, adapted for Ollama).
                # semantic_contextual_late— all three: LLM context + late context window
                #                          + semantic splits. Highest quality, slowest.
                enriched_parts = []
                if method_used in ("contextual", "semantic_contextual_late"):
                    ctx = _generate_chunk_context(full_text, text)
                    if ctx:
                        enriched_parts.append(f"[Context: {ctx}]")
                    elif topics_str:
                        enriched_parts.append(f"[Topics: {topics_str}]")
                elif topics_str:
                    enriched_parts.append(f"[Topics: {topics_str}]")

                if method_used in ("late_chunking", "semantic_contextual_late"):
                    late_ctx = _build_late_chunk_context(splits, idx)
                    if late_ctx:
                        enriched_parts.append(f"[LateCtx: {late_ctx}]")

                if heading_ctx:
                    enriched_parts.append(f"[Section: {heading_ctx}]")
                enriched_parts.append(text)
                enriched_text = "\n".join(enriched_parts)

                docs_payload.append(enriched_text)
                chunk_meta: dict[str, str] = {
                    "project": project_name,
                        "source": path.name,
                    "theme": theme_name,
                        "chunking_method": method_used,
                        "file_path": str(path),
                    "source_fingerprint": source_fingerprint,
                    "chunk_fingerprint": chunk_fingerprint,
                        "section": heading_ctx[:200] if heading_ctx else "",
                        "topics": topics_str[:200],
                        "date_added": date_added,
                        "page_number": page_number,
                        # Harvested PDF metadata for APA citations
                        "doc_title": doc_meta.get("doc_title", "")[:300],
                        "doc_authors": doc_meta.get("doc_authors", "")[:300],
                        "doc_year": doc_meta.get("doc_year", ""),
                        # Verbatim source text (no LLM enrichment) — used by rag_retrieve
                        # to display grounded, quotable content instead of the embedding-enriched text.
                        "chunk_source_text": text[:4000],
                }
                # Merge any extra metadata from literature search (doi, journal, etc.)
                if extra_meta:
                    for k, v in extra_meta.items():
                        chunk_meta[k] = str(v)[:300]
                metas_payload.append(chunk_meta)
                ids_payload.append(
                    _build_chunk_id(
                        project=project_name,
                        theme=theme_name,
                        source_fingerprint=source_fingerprint,
                        chunk_fingerprint=chunk_fingerprint,
                    )
                )

            if not docs_payload:
                failures.append(f"{path} -> chunks were empty after text cleanup")
                continue

            if not (docs_payload and metas_payload and ids_payload):
                failures.append(f"{path} -> skipped due to empty upsert payload")
                continue

            if not (len(docs_payload) == len(metas_payload) == len(ids_payload)):
                failures.append(f"{path} -> skipped due to inconsistent chunk payload sizes")
                continue

            table_docs_payload: list[str] = []
            table_metas_payload: list[dict[str, str]] = []
            table_ids_payload: list[str] = []
            image_docs_payload: list[str] = []
            image_metas_payload: list[dict[str, str]] = []
            image_ids_payload: list[str] = []
            video_frame_docs_payload: list[str] = []
            video_frame_metas_payload: list[dict[str, str]] = []
            video_frame_ids_payload: list[str] = []
            if path.suffix.lower() == ".pdf":
                table_docs_payload, table_metas_payload, table_ids_payload = _extract_pdf_table_records(
                    path,
                    docs,
                    project_name=project_name,
                    theme_name=theme_name,
                    source_fingerprint=source_fingerprint,
                    date_added=date_added,
                    doc_meta=doc_meta,
                    extra_meta=extra_meta,
                )
                if table_docs_payload:
                    table_docs_payload, table_metas_payload, table_ids_payload = _dedupe_multimodal_payloads(
                        docs_payload,
                        metas_payload,
                        table_docs_payload,
                        table_metas_payload,
                        table_ids_payload,
                    )
                    docs_payload.extend(table_docs_payload)
                    metas_payload.extend(table_metas_payload)
                    ids_payload.extend(table_ids_payload)

                image_docs_payload, image_metas_payload, image_ids_payload = _extract_pdf_image_records(
                    path,
                    docs,
                    project_name=project_name,
                    theme_name=theme_name,
                    source_fingerprint=source_fingerprint,
                    date_added=date_added,
                    doc_meta=doc_meta,
                    extra_meta=extra_meta,
                )
                if image_docs_payload:
                    image_docs_payload, image_metas_payload, image_ids_payload = _dedupe_multimodal_payloads(
                        docs_payload,
                        metas_payload,
                        image_docs_payload,
                        image_metas_payload,
                        image_ids_payload,
                    )
                    docs_payload.extend(image_docs_payload)
                    metas_payload.extend(image_metas_payload)
                    ids_payload.extend(image_ids_payload)
            elif path.suffix.lower() in IMAGE_FILE_SUFFIXES:
                image_docs_payload, image_metas_payload, image_ids_payload = _extract_standalone_image_records(
                    path,
                    project_name=project_name,
                    theme_name=theme_name,
                    source_fingerprint=source_fingerprint,
                    date_added=date_added,
                    doc_meta=doc_meta,
                    extra_meta=extra_meta,
                )
                if image_docs_payload:
                    image_docs_payload, image_metas_payload, image_ids_payload = _dedupe_multimodal_payloads(
                        docs_payload,
                        metas_payload,
                        image_docs_payload,
                        image_metas_payload,
                        image_ids_payload,
                    )
                    docs_payload.extend(image_docs_payload)
                    metas_payload.extend(image_metas_payload)
                    ids_payload.extend(image_ids_payload)
            elif path.suffix.lower() in VIDEO_FILE_SUFFIXES:
                video_frame_docs_payload, video_frame_metas_payload, video_frame_ids_payload = _extract_video_frame_records(
                    path,
                    project_name=project_name,
                    theme_name=theme_name,
                    source_fingerprint=source_fingerprint,
                    date_added=date_added,
                    doc_meta=doc_meta,
                    extra_meta=extra_meta,
                )
                if video_frame_docs_payload:
                    video_frame_docs_payload, video_frame_metas_payload, video_frame_ids_payload = _dedupe_multimodal_payloads(
                        docs_payload,
                        metas_payload,
                        video_frame_docs_payload,
                        video_frame_metas_payload,
                        video_frame_ids_payload,
                    )
                    docs_payload.extend(video_frame_docs_payload)
                    metas_payload.extend(video_frame_metas_payload)
                    ids_payload.extend(video_frame_ids_payload)

            stale_chunks_deleted += _delete_stale_chunks(
                collection=collection,
                source=path.name,
                theme=theme_name,
                source_fingerprint=source_fingerprint,
            )

            try:
                collection.upsert(
                    documents=docs_payload,
                    metadatas=metas_payload,
                    ids=ids_payload,
                )
            except ValueError as ve:
                # Chroma can raise this for malformed/empty payloads. Convert to file-level
                # failure and continue processing other files.
                failures.append(f"{path} -> upsert validation error: {ve}")
                continue

            loaded_files += 1
            added_chunks += len(docs_payload)
            _progress(
                "ingest", file_idx + 1, total_files,
                f"[{file_idx + 1}/{total_files}] {path.name} — {len(docs_payload)} chunks",
            )
        except Exception as e:
            failures.append(f"{path} -> {e}")

    return {
        "loaded_files": loaded_files,
        "total_files": len(paths),
        "added_chunks": added_chunks,
        "stopped": _stopped(),
        "project": project_name,
        "theme": theme_name,
        "method_used": method_used,
        "stale_chunks_deleted": stale_chunks_deleted,
        "failures": failures,
    }


def _delete_stale_chunks(
    collection: Any,
    *,
    source: str,
    theme: str,
    source_fingerprint: str,
) -> int:
    """Delete only stale chunks for the same source within the same project+theme."""
    try:
        count = collection.count()
    except Exception:
        return 0

    if count == 0:
        return 0

    data = _safe_collection_snapshot(collection, include=["metadatas"])
    if not data:
        return 0

    ids = data.get("ids", []) or []
    metadatas = data.get("metadatas", []) or []
    stale_ids = [
        doc_id
        for doc_id, meta in zip(ids, metadatas)
        if (meta or {}).get("source") == source
        and (meta or {}).get("theme") == theme
        and (meta or {}).get("source_fingerprint") == source_fingerprint
    ]
    if not stale_ids:
        return 0

    collection.delete(ids=stale_ids)
    return len(stale_ids)


def _safe_collection_snapshot(collection: Any, include: list[str] | None = None) -> dict[str, Any]:
    """Best-effort collection read compatible with real Chroma and light test doubles."""
    include = include or ["metadatas"]
    try:
        data = collection.get(include=include)
    except Exception:
        data = None

    ids = list((data or {}).get("ids", []) or [])
    needs_ids = not ids
    needs_docs = "documents" in include and not list((data or {}).get("documents", []) or [])
    needs_meta = "metadatas" in include and not list((data or {}).get("metadatas", []) or [])

    if needs_ids or needs_docs or needs_meta:
        try:
            peek_count = collection.count()
            peek_data = collection.peek(peek_count)
        except Exception:
            peek_data = {}
        merged = dict(peek_data or {})
        merged.update({k: v for k, v in (data or {}).items() if v})
        data = merged

    return data or {}


def _list_rag_collections() -> list:
    """List all ChromaDB collections for the current embed model across all projects (B2)."""
    try:
        import chromadb

        embed_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
        model_hash = hashlib.sha1(embed_model.encode("utf-8")).hexdigest()[:8]
        client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))
        return [c for c in client.list_collections() if c.name.endswith(f"__{model_hash}")]
    except Exception:
        return []


def get_rag_index_summary(project: str = "Default") -> dict[str, Any]:
    """Return indexed files, themes, and counts for a single project.

    The returned dict always contains ``project``, ``total_chunks``, and
    ``files``.  If the collection exists but is unreadable (e.g. index
    corruption) the dict will also contain an ``error`` key with a
    human-readable message and ``total_chunks`` set to the raw count so
    the caller knows data *was* written.
    """
    project_name = _normalize_project(project)
    summary: dict[str, Any] = {
        "project": project_name,
        "total_chunks": 0,
        "files": {},
    }

    try:
        collection = _get_vector_collection(project_name)
        total = collection.count()
        if total == 0:
            return summary
        data = _safe_collection_snapshot(collection, include=["metadatas", "documents"])
        if not data:
            return summary
    except Exception:
        # Fallback: read directly from Chroma without requiring embedding services.
        try:
            import chromadb
            collection_name = _get_rag_collection_name(project_name)
            client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))
            try:
                collection = client.get_collection(name=collection_name)
            except Exception:
                return summary
            total = collection.count()
            if total == 0:
                return summary
            data = _safe_collection_snapshot(collection, include=["metadatas", "documents"])
            if not data:
                summary["total_chunks"] = total
                summary["error"] = (
                    f"Collection '{collection_name}' reports {total} chunks but cannot be read. "
                    "The index may be corrupted — delete and re-ingest."
                )
                return summary
        except Exception:
            return summary

    documents = data.get("documents", [])
    for idx, meta in enumerate(data.get("metadatas", [])):
        meta = meta or {}
        document_text = documents[idx] if idx < len(documents) else ""
        source = meta.get("source", "unknown")
        theme = meta.get("theme") or "Unspecified"
        date_added = meta.get("date_added", "")
        summary["total_chunks"] += 1
        summary["files"].setdefault(
            source,
            {
                "chunks": 0,
                "themes": set(),
                "theme_counts": {},
                "chunking_methods": set(),
                "modalities": {},
                "table_extraction_methods": {},
                "vision_captioned_chunks": 0,
                "date_added": date_added,
            },
        )
        summary["files"][source]["chunks"] += 1
        summary["files"][source]["themes"].add(theme)
        chunking_method = _infer_chunking_method_from_document(document_text, meta)
        if chunking_method:
            summary["files"][source]["chunking_methods"].add(chunking_method)
        modality = str(meta.get("modality", "text") or "text").lower()
        summary["files"][source]["modalities"][modality] = (
            summary["files"][source]["modalities"].get(modality, 0) + 1
        )
        table_extraction_method = str(meta.get("table_extraction_method", "") or "").strip().lower()
        if table_extraction_method:
            summary["files"][source]["table_extraction_methods"][table_extraction_method] = (
                summary["files"][source]["table_extraction_methods"].get(table_extraction_method, 0) + 1
            )
        if str(meta.get("vision_caption_source", "") or "").strip():
            summary["files"][source]["vision_captioned_chunks"] += 1
        summary["files"][source]["theme_counts"][theme] = (
            summary["files"][source]["theme_counts"].get(theme, 0) + 1
        )
        if date_added and (
            not summary["files"][source]["date_added"]
            or date_added < summary["files"][source]["date_added"]
        ):
            summary["files"][source]["date_added"] = date_added
    return summary


def get_rag_available_themes(project: str = "Default") -> list[str]:
    """Get sorted list of unique themes for a single project."""
    summary = get_rag_index_summary(project=project)
    themes: set[str] = set()
    for info in summary.get("files", {}).values():
        themes.update(info.get("themes", set()))
    return sorted(themes)


def get_rag_projects() -> list[str]:
    """Get sorted list of known RAG projects that have at least one chunk.

    Derives project names from collection names (fast — no data reads) and
    verifies each has a non-zero count.
    """
    projects: set[str] = set()
    try:
        import chromadb
        embed_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
        model_hash = hashlib.sha1(embed_model.encode("utf-8")).hexdigest()[:8]
        suffix = f"__{model_hash}"
        client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))
        for c in client.list_collections():
            if not c.name.endswith(suffix):
                continue
            # Extract project slug from "rag__{slug}__{hash}"
            prefix = "rag__"
            if not c.name.startswith(prefix):
                continue
            slug = c.name[len(prefix):-len(suffix)]
            if not slug:
                continue
            try:
                if c.count() == 0:
                    continue
            except Exception:
                continue
            # Recover human-readable name from first chunk's metadata
            try:
                sample = c.get(include=["metadatas"], limit=1)
                name = ((sample.get("metadatas") or [{}])[0] or {}).get("project", "").strip()
                projects.add(name if name else slug.replace("_", " ").title())
            except Exception:
                projects.add(slug.replace("_", " ").title())
    except Exception:
        pass
    return sorted(projects)


def delete_rag_documents(
    project: str = "Default",
    source: str = "",
    theme: str = "",
    allow_delete_all: bool = False,
) -> int:
    """Delete matching chunks by source/theme within a single project."""
    project_name = _normalize_project(project)
    source_name = (source or "").strip()
    theme_name = (theme or "").strip()

    if not allow_delete_all and not source_name and not theme_name:
        return 0

    try:
        collection = _get_vector_collection(project_name)
        count = collection.count()
    except Exception:
        return 0

    if count == 0:
        return 0

    data = _safe_collection_snapshot(collection, include=["metadatas"])
    if not data:
        return 0
    ids = data.get("ids", []) or []
    metadatas = data.get("metadatas", []) or []
    to_delete: list[str] = []
    for doc_id, meta in zip(ids, metadatas):
        meta = meta or {}
        source_match = not source_name or meta.get("source") == source_name
        theme_match = not theme_name or meta.get("theme") == theme_name
        if source_match and theme_match:
            to_delete.append(doc_id)

    if not to_delete:
        return 0

    collection.delete(ids=to_delete)
    return len(to_delete)


_vector_collection_cache: dict[str, Any] = {}
_vector_collection_model: str = ""


def _get_vector_collection(project: str = "Default"):
    """Return the ChromaDB collection for a project; rebuilds cache on embed model change (B2)."""
    global _vector_collection_cache, _vector_collection_model
    current_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
    if current_model != _vector_collection_model:
        _vector_collection_cache.clear()
        _vector_collection_model = current_model
    if project not in _vector_collection_cache:
        _vector_collection_cache[project] = _create_vector_collection(project)
    return _vector_collection_cache[project]


def _create_vector_collection(project: str = "Default"):
    try:
        import chromadb
    except Exception as e:
        raise RuntimeError(
            "RAG dependencies missing. Install: chromadb"
        ) from e

    embedding_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
    llama_server_base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
    llama_cpp_n_ctx = int(os.getenv("RAG_LLAMA_CPP_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
    llama_cpp_n_batch = int(os.getenv("RAG_LLAMA_CPP_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
    llama_cpp_n_gpu_layers = int(os.getenv("RAG_LLAMA_CPP_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
    llama_cpp_flash_attn = os.getenv(
        "RAG_LLAMA_CPP_FLASH_ATTN",
        os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
    ).strip().lower() in {"1", "true", "yes", "on"}
    collection_name = _get_rag_collection_name(project)

    ef = _get_rag_embedding_function(
        embedding_model,
        llama_server_base_url,
        llama_cpp_n_ctx,
        llama_cpp_n_batch,
        llama_cpp_n_gpu_layers,
        llama_cpp_flash_attn,
    )
    client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))
    return client.get_or_create_collection(name=collection_name, embedding_function=ef)


def _harvest_docx_metadata(file_path: Path) -> dict[str, str]:
    """Extract document metadata from a .docx file using python-docx (B5).

    Returns a dict with keys: doc_title, doc_authors, doc_year.
    Falls back to empty strings when metadata is absent or python-docx is unavailable.
    """
    meta: dict[str, str] = {"doc_title": "", "doc_authors": "", "doc_year": ""}
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        props = doc.core_properties
        meta["doc_title"] = (props.title or "").strip()
        meta["doc_authors"] = (props.author or "").strip()
        created = props.created
        modified = props.modified
        if created:
            meta["doc_year"] = str(created.year)
        elif modified:
            meta["doc_year"] = str(modified.year)
    except Exception:
        pass
    return meta


def _load_documents(file_path: Path) -> tuple[list, dict[str, str]]:
    """Load a document and return (pages, doc_metadata).

    For PDFs, harvests real document metadata (title, authors, year) via
    PyMuPDF and injects page-break sentinels so chunks respect page boundaries.
    Returns a tuple: (langchain Document list, harvested metadata dict).
    """
    suffix = file_path.suffix.lower()
    if suffix in IMAGE_FILE_SUFFIXES:
        try:
            from langchain_core.documents import Document as LCDocument
        except ImportError:
            from langchain.schema import Document as LCDocument

        return [
            LCDocument(
                page_content=f"Standalone image file: {file_path.name}",
                metadata={"page_number": "1", "media_kind": "image"},
            )
        ], {}
    if suffix in AUDIO_FILE_SUFFIXES | VIDEO_FILE_SUFFIXES:
        try:
            from langchain_core.documents import Document as LCDocument
        except ImportError:
            from langchain.schema import Document as LCDocument

        transcript = _transcribe_media_file(file_path)
        media_kind = "video" if suffix in VIDEO_FILE_SUFFIXES else "audio"
        return [
            LCDocument(
                page_content=transcript,
                metadata={"page_number": "", "media_kind": media_kind},
            )
        ], {}
    if suffix == ".pdf":
        from langchain_community.document_loaders import PyMuPDFLoader
        try:
            from langchain_core.documents import Document as LCDocument
        except ImportError:
            # Backward compatibility for older LangChain releases.
            from langchain.schema import Document as LCDocument

        pages = PyMuPDFLoader(str(file_path)).load()
        doc_meta = _harvest_pdf_metadata(file_path)

        # Inject page number into each page's metadata so chunks inherit it
        enriched = []
        for page_doc in pages:
            page_num = page_doc.metadata.get("page", "")  # PyMuPDF uses "page" (0-based)
            display_page = int(page_num) + 1 if str(page_num).isdigit() else page_num
            new_meta = {**page_doc.metadata, "page_number": str(display_page)}
            enriched.append(LCDocument(page_content=page_doc.page_content, metadata=new_meta))
        return enriched, doc_meta

    if suffix == ".docx":
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader

        doc_meta = _harvest_docx_metadata(file_path)
        return UnstructuredWordDocumentLoader(str(file_path)).load(), doc_meta
    if suffix in [".doc", ".odt"]:
        from langchain_community.document_loaders import UnstructuredFileLoader
        return UnstructuredFileLoader(str(file_path)).load(), {}
    if suffix in [".xlsx", ".xls"]:
        try:
            from langchain_community.document_loaders import UnstructuredExcelLoader

            return UnstructuredExcelLoader(str(file_path)).load(), {}
        except ImportError:
            from langchain_community.document_loaders import UnstructuredFileLoader

            return UnstructuredFileLoader(str(file_path)).load(), {}
    if suffix in [".pptx", ".ppt"]:
        try:
            from langchain_community.document_loaders import UnstructuredPowerPointLoader

            return UnstructuredPowerPointLoader(str(file_path)).load(), {}
        except ImportError:
            from langchain_community.document_loaders import UnstructuredFileLoader

            return UnstructuredFileLoader(str(file_path)).load(), {}
    if suffix in [".txt", ".md", ".csv", ".json"]:
        from langchain_community.document_loaders import TextLoader
        return TextLoader(str(file_path), encoding="utf-8").load(), {}

    raise ValueError(
        f"Unsupported file extension: {suffix}. Supported: .pdf, .docx, .doc, .odt, .pptx, .ppt, .xlsx, .xls, .txt, .md, .csv, .json, .png, .jpg, .jpeg, .webp, .bmp, .gif, .tif, .tiff, .mp3, .wav, .m4a, .aac, .flac, .ogg, .opus, .mp4, .mov, .mkv, .avi, .webm"
    )


def _meta_matches_filter(meta: dict[str, Any], where_filter: dict) -> bool:
    """Check whether metadata satisfies a simple ChromaDB where filter (handles $and/$or/$eq/$in)."""
    if "$and" in where_filter:
        return all(_meta_matches_filter(meta, clause) for clause in where_filter["$and"])
    if "$or" in where_filter:
        return any(_meta_matches_filter(meta, clause) for clause in where_filter["$or"])
    for key, condition in where_filter.items():
        if key.startswith("$"):
            continue
        val = meta.get(key, "")
        if isinstance(condition, dict):
            if "$eq" in condition and val != condition["$eq"]:
                return False
            if "$in" in condition and val not in condition["$in"]:
                return False
        elif val != condition:
            return False
    return True


def _bm25_search(
    query: str,
    documents: list[str],
    metadatas: list[dict[str, Any]],
    ids: list[str],
    top_k: int,
    prebuilt_bm25: Any = None,
) -> list[tuple[str, dict[str, Any], str]]:
    """BM25 keyword search over a list of documents (B1).

    Uses rank_bm25 (BM25Okapi) when available; falls back to sklearn TF-IDF cosine.
    Accepts a prebuilt BM25 index to avoid re-tokenizing on every query.
    Returns at most top_k results with positive relevance scores.
    """
    if not documents:
        return []
    try:
        import numpy as np

        try:
            from rank_bm25 import BM25Okapi

            if prebuilt_bm25 is not None:
                bm25 = prebuilt_bm25
            else:
                tokenized = [_tokenize_for_bm25(doc) for doc in documents]
                bm25 = BM25Okapi(tokenized)
            scores = bm25.get_scores(_tokenize_for_bm25(query))
        except ImportError:
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.metrics.pairwise import cosine_similarity

            vec = TfidfVectorizer(max_features=10000, stop_words="english")
            mat = vec.fit_transform(documents)
            q_vec = vec.transform([query])
            scores = cosine_similarity(q_vec, mat)[0]

        ranked_idx = np.argsort(scores)[::-1][:top_k]
        return [
            (documents[int(i)], metadatas[int(i)], ids[int(i)])
            for i in ranked_idx
            if scores[int(i)] > 0
        ]
    except Exception:
        return []


def _reciprocal_rank_fusion(
    result_lists: list[list[tuple[str, dict[str, Any], str]]],
    k: int = 60,
) -> list[tuple[str, dict[str, Any], str]]:
    """Merge multiple ranked lists via Reciprocal Rank Fusion (RRF) (B1).

    RRF score = Σ 1/(k + rank_i) across all lists.
    k=60 is the standard dampening constant (Robertson et al.).
    Deduplicates by doc_id — first occurrence wins for the stored doc/meta.
    """
    scores: dict[str, float] = {}
    doc_map: dict[str, tuple[str, dict[str, Any], str]] = {}
    for result_list in result_lists:
        for rank, (doc, meta, doc_id) in enumerate(result_list):
            scores.setdefault(doc_id, 0.0)
            scores[doc_id] += 1.0 / (k + rank + 1)
            doc_map.setdefault(doc_id, (doc, meta, doc_id))
    sorted_ids = sorted(scores, key=lambda x: scores[x], reverse=True)
    return [doc_map[d] for d in sorted_ids]


def _query_all_matching_collections(
    query: str,
    n_results: int,
    where_filter: dict | None = None,
) -> tuple[list[str], list[dict[str, Any]], list[str]]:
    """Query every ChromaDB collection matching the current embed model and merge results.

    Uses ThreadPoolExecutor to query collections in parallel for lower latency.
    Deduplicates by chunk ID; first-seen instance wins.
    """
    try:
        import chromadb

        embedding_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
        llama_server_base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
        llama_cpp_n_ctx = int(os.getenv("RAG_LLAMA_CPP_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
        llama_cpp_n_batch = int(os.getenv("RAG_LLAMA_CPP_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
        llama_cpp_n_gpu_layers = int(os.getenv("RAG_LLAMA_CPP_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
        llama_cpp_flash_attn = os.getenv(
            "RAG_LLAMA_CPP_FLASH_ATTN",
            os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
        ).strip().lower() in {"1", "true", "yes", "on"}
        model_hash = hashlib.sha1(embedding_model.encode("utf-8")).hexdigest()[:8]
        ef = _get_rag_embedding_function(
            embedding_model,
            llama_server_base_url,
            llama_cpp_n_ctx,
            llama_cpp_n_batch,
            llama_cpp_n_gpu_layers,
            llama_cpp_flash_attn,
        )
        client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))

        matching_colls = [
            c for c in client.list_collections()
            if (c.name if hasattr(c, "name") else str(c)).endswith(f"__{model_hash}")
        ]
        if not matching_colls:
            return [], [], []

        def _query_single(coll_meta: Any) -> list[tuple[str, dict, str]]:
            coll_name = coll_meta.name if hasattr(coll_meta, "name") else str(coll_meta)
            try:
                coll = client.get_collection(coll_name, embedding_function=ef)
                count = coll.count()
                if count == 0:
                    return []
                result = coll.query(
                    query_texts=[query],
                    n_results=min(n_results, count),
                    where=where_filter,
                )
                return [
                    (str(doc), meta or {}, str(doc_id))
                    for doc, meta, doc_id in zip(
                        result.get("documents", [[]])[0],
                        result.get("metadatas", [[]])[0],
                        result.get("ids", [[]])[0],
                    )
                ]
            except Exception:
                return []

        all_docs: list[str] = []
        all_metas: list[dict[str, Any]] = []
        all_ids: list[str] = []
        seen_ids: set[str] = set()

        # Parallel query across collections
        max_workers = min(len(matching_colls), 4)
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(_query_single, c): c for c in matching_colls}
            for future in as_completed(futures):
                for doc, meta, doc_id in future.result():
                    if doc_id not in seen_ids:
                        seen_ids.add(doc_id)
                        all_docs.append(doc)
                        all_metas.append(meta)
                        all_ids.append(doc_id)

        return all_docs, all_metas, all_ids
    except Exception:
        return [], [], []


def _score_ranked_candidates(
    encoder_model: Any,
    query: str,
    doc_triples: list[tuple[str, dict[str, Any], str]],
) -> list[tuple[str, dict[str, Any], str, float]]:
    """Score candidate chunks with the cross encoder and keep score metadata."""
    texts = [doc for doc, _, _ in doc_triples]
    ranks = encoder_model.rank(query, texts, top_k=len(texts))
    intent_boosts = _get_modality_intent_boosts(query)
    scored = [
        (
            doc_triples[rank["corpus_id"]][0],
            doc_triples[rank["corpus_id"]][1],
            doc_triples[rank["corpus_id"]][2],
            float(rank.get("score", 1.0))
            + _score_modality_bonus(doc_triples[rank["corpus_id"]][1], intent_boosts),
        )
        for rank in ranks
    ]
    return sorted(scored, key=lambda item: item[3], reverse=True)


def _get_modality_intent_boosts(query: str) -> dict[str, float]:
    """Infer lightweight modality preferences from the query."""
    text = (query or "").lower()
    boosts: dict[str, float] = {}

    image_markers = {
        "image", "images", "figure", "figures", "diagram", "diagrams", "screenshot",
        "screenshots", "photo", "photos", "visual", "visuals", "chart", "charts",
        "graph", "graphs", "picture", "pictures",
    }
    table_markers = {
        "table", "tables", "tabular", "row", "rows", "column", "columns",
        "compare values", "matrix", "spreadsheet",
    }

    if any(marker in text for marker in image_markers):
        boosts["image"] = 0.12
    if any(marker in text for marker in table_markers):
        boosts["table"] = 0.12

    return boosts


def _score_modality_bonus(meta: dict[str, Any], intent_boosts: dict[str, float]) -> float:
    """Return a small additive rerank bonus for modality-aligned chunks."""
    modality = str((meta or {}).get("modality", "text") or "text").lower()
    return float(intent_boosts.get(modality, 0.0))


def _round_robin_file_candidates(
    selected_by_file: list[tuple[str, list[tuple[str, dict[str, Any], str, float]]]],
    top_k: int,
) -> list[tuple[str, dict[str, Any], str]]:
    """Interleave best chunks across selected files so one file cannot dominate."""
    chosen: list[tuple[str, dict[str, Any], str]] = []
    position = 0
    while len(chosen) < top_k:
        added_in_round = False
        for _, candidates in selected_by_file:
            if position < len(candidates):
                doc, meta, doc_id, _score = candidates[position]
                chosen.append((doc, meta, doc_id))
                added_in_round = True
                if len(chosen) >= top_k:
                    break
        if not added_in_round:
            break
        position += 1
    return chosen


def _rerank_chunks(
    query: str,
    documents: list[str],
    metadatas: list[dict[str, Any]],
    ids: list[str],
    top_k: int,
    mode: str,
    max_files: int,
    min_score: float = 0.0,
    return_details: bool = False,
) -> Any:
    """Rerank chunks using CrossEncoder. Returns (doc, meta, id) triples.

    Uses adaptive percentile-based score filtering: instead of a fixed min_score,
    computes the 25th percentile of all CrossEncoder scores and uses the higher
    of (min_score, percentile_threshold). This adapts automatically to different
    query/corpus combinations.

    Args:
        min_score: Floor threshold — chunks below this are always discarded.
                   Defaults to 0.0. Set via RAG_MIN_RERANK_SCORE env var.
    """
    diagnostics: dict[str, Any] = {
        "candidate_count": len(documents),
        "rerank_count": 0,
        "filtered_count": 0,
        "selected_files": [],
        "file_scores": {},
        "reranker": "cross-encoder",
    }

    try:
        encoder_model = _get_cross_encoder()
    except Exception:
        # Fallback without reranker — no score filtering possible
        fallback = list(zip(documents, metadatas, ids))
        diagnostics.update(
            {
                "reranker": "fallback",
                "rerank_count": len(fallback),
                "filtered_count": len(fallback[:top_k]),
            }
        )
        results = fallback[:top_k]
        return (results, diagnostics) if return_details else results

    if mode in ("MMR", "Top-K Per File"):
        grouped: dict[str, list[tuple[str, dict[str, Any], str]]] = {}
        for doc, meta, doc_id in zip(documents, metadatas, ids):
            source = (meta or {}).get("source", "unknown")
            grouped.setdefault(source, []).append((doc, meta or {}, doc_id))

        rerank_count = 0
        filtered_count = 0
        file_scores: dict[str, float] = {}
        all_scores_for_percentile: list[float] = []
        scored_by_file: dict[str, list[tuple[str, dict[str, Any], str, float]]] = {}
        for source, doc_triples in grouped.items():
            scored = _score_ranked_candidates(encoder_model, query, doc_triples)
            rerank_count += len(scored)
            if scored:
                file_scores[source] = max(score for *_rest, score in scored)
                all_scores_for_percentile.extend(score for *_rest, score in scored)
            scored_by_file[source] = scored

        # Adaptive threshold: use the higher of min_score and 25th percentile
        adaptive_threshold = min_score
        if all_scores_for_percentile:
            percentile = float(os.getenv("RAG_RERANK_PERCENTILE", "25"))
            sorted_scores = sorted(all_scores_for_percentile)
            p_idx = int(len(sorted_scores) * percentile / 100)
            p_idx = min(p_idx, len(sorted_scores) - 1)
            adaptive_threshold = max(min_score, sorted_scores[p_idx])

        filtered_by_file: dict[str, list[tuple[str, dict[str, Any], str, float]]] = {}
        for source, scored in scored_by_file.items():
            filtered = [item for item in scored if item[3] >= adaptive_threshold]
            filtered_count += len(filtered)
            if filtered:
                filtered_by_file[source] = filtered

        selected_sources = [
            source
            for source, _score in sorted(
                file_scores.items(),
                key=lambda item: item[1],
                reverse=True,
            )
            if source in filtered_by_file
        ][:max_files]
        diagnostics.update(
            {
                "rerank_count": rerank_count,
                "filtered_count": filtered_count,
                "selected_files": selected_sources,
                "file_scores": {
                    source: round(file_scores[source], 6)
                    for source in selected_sources
                },
            }
        )

        selected_by_file = [
            (source, filtered_by_file[source])
            for source in selected_sources
        ]
        selected: list[tuple[str, dict[str, Any], str, float]] = [
            candidate
            for _source, candidates in selected_by_file
            for candidate in candidates
        ]

        # Semantic fallback: if CrossEncoder filtered everything, return top-K
        # by original embedding similarity (ChromaDB order) instead of nothing.
        if not selected:
            fallback = list(zip(documents, metadatas, ids))[:top_k]
            diagnostics.update({
                "reranker": "semantic-fallback",
                "filtered_count": len(fallback),
            })
            return (fallback, diagnostics) if return_details else fallback

        if mode == "MMR":
            results = _apply_mmr(
                sorted(selected, key=lambda item: item[3], reverse=True),
                top_k=top_k,
            )
            return (results, diagnostics) if return_details else results
        results = _round_robin_file_candidates(selected_by_file, top_k=top_k)
        return (results, diagnostics) if return_details else results

    scored = _score_ranked_candidates(
        encoder_model,
        query,
        list(zip(documents, metadatas, ids)),
    )
    # Adaptive threshold: use the higher of min_score and 25th percentile
    adaptive_threshold = min_score
    if scored:
        percentile = float(os.getenv("RAG_RERANK_PERCENTILE", "25"))
        sorted_scores = sorted(score for *_rest, score in scored)
        p_idx = int(len(sorted_scores) * percentile / 100)
        p_idx = min(p_idx, len(sorted_scores) - 1)
        adaptive_threshold = max(min_score, sorted_scores[p_idx])
    filtered = [
        (doc, meta, doc_id)
        for doc, meta, doc_id, score in scored
        if score >= adaptive_threshold
    ]
    # Semantic fallback: if CrossEncoder filtered everything, return top-K
    # by original embedding similarity (ChromaDB order) instead of nothing.
    if not filtered and scored:
        fallback = [(doc, meta, doc_id) for doc, meta, doc_id, _score in scored][:top_k]
        diagnostics.update({
            "reranker": "semantic-fallback",
            "rerank_count": len(scored),
            "filtered_count": len(fallback),
        })
        return (fallback, diagnostics) if return_details else fallback
    diagnostics.update(
        {
            "rerank_count": len(scored),
            "filtered_count": len(filtered),
        }
    )
    results = filtered[:top_k]
    return (results, diagnostics) if return_details else results


def _apply_mmr(
    candidates: list[tuple[str, dict[str, Any], str, float]],
    top_k: int,
    lambda_param: float = 0.5,
) -> list[tuple[str, dict[str, Any], str]]:
    """Maximal Marginal Relevance post-processing to suppress near-duplicate chunks.

    Iteratively selects the candidate that maximises:
        lambda * relevance_score - (1 - lambda) * max_similarity_to_already_selected

    Relevance uses real CrossEncoder scores (passed in as the 4th tuple element).
    Similarity uses real embedding vectors from ChromaDB when available, falling
    back to TF-IDF cosine distance.

    lambda_param=0 → max diversity, lambda_param=1 → max relevance.
    Default 0.5 = balanced. Override via RAG_MMR_LAMBDA env var.
    """
    if len(candidates) <= top_k:
        return [(doc, meta, doc_id) for doc, meta, doc_id, _ in candidates]

    lambda_val = float(os.getenv("RAG_MMR_LAMBDA", str(lambda_param)))

    try:
        import numpy as np
        from sklearn.metrics.pairwise import cosine_similarity as cos_sim

        texts = [doc for doc, _, _, _ in candidates]
        ce_scores = [score for _, _, _, score in candidates]

        # Try to get real embeddings from ChromaDB for better diversity estimation
        sim_matrix = None
        try:
            chunk_ids = [doc_id for _, _, doc_id, _ in candidates]
            import chromadb
            embedding_model = os.getenv("RAG_EMBED_MODEL", "llama_cpp:").strip()
            llama_server_base_url = os.getenv("LLAMA_SERVER_BASE_URL", "http://localhost:8080/v1").rstrip("/")
            llama_cpp_n_ctx = int(os.getenv("RAG_LLAMA_CPP_N_CTX", os.getenv("LLAMA_CPP_N_CTX", "8192")))
            llama_cpp_n_batch = int(os.getenv("RAG_LLAMA_CPP_N_BATCH", os.getenv("LLAMA_CPP_N_BATCH", "256")))
            llama_cpp_n_gpu_layers = int(os.getenv("RAG_LLAMA_CPP_N_GPU_LAYERS", os.getenv("LLAMA_CPP_N_GPU_LAYERS", "-1")))
            llama_cpp_flash_attn = os.getenv(
                "RAG_LLAMA_CPP_FLASH_ATTN",
                os.getenv("LLAMA_CPP_FLASH_ATTN", "true"),
            ).strip().lower() in {"1", "true", "yes", "on"}
            ef = _get_rag_embedding_function(
                embedding_model,
                llama_server_base_url,
                llama_cpp_n_ctx,
                llama_cpp_n_batch,
                llama_cpp_n_gpu_layers,
                llama_cpp_flash_attn,
            )
            client = chromadb.PersistentClient(path=str(RAG_CHROMA_DIR))
            model_hash = hashlib.sha1(embedding_model.encode("utf-8")).hexdigest()[:8]
            # Gather embeddings from matching collections
            embeddings_map: dict[str, list[float]] = {}
            for coll_meta in client.list_collections():
                coll_name = coll_meta.name if hasattr(coll_meta, "name") else str(coll_meta)
                if not coll_name.endswith(f"__{model_hash}"):
                    continue
                needed = [cid for cid in chunk_ids if cid not in embeddings_map]
                if not needed:
                    break
                try:
                    coll = client.get_collection(coll_name, embedding_function=ef)
                    result = coll.get(ids=needed, include=["embeddings"])
                    for rid, emb in zip(result.get("ids", []), result.get("embeddings", [])):
                        if emb is not None:
                            embeddings_map[rid] = emb
                except Exception:
                    continue
            if len(embeddings_map) == len(chunk_ids):
                emb_matrix = np.array([embeddings_map[cid] for cid in chunk_ids])
                sim_matrix = cos_sim(emb_matrix)
        except Exception:
            pass

        # Fallback to TF-IDF if embeddings unavailable
        if sim_matrix is None:
            from sklearn.feature_extraction.text import TfidfVectorizer
            vectorizer = TfidfVectorizer(max_features=5000, stop_words="english")
            tfidf_matrix = vectorizer.fit_transform(texts)
            sim_matrix = cos_sim(tfidf_matrix).toarray() if hasattr(cos_sim(tfidf_matrix), 'toarray') else cos_sim(tfidf_matrix)
            # Re-compute properly
            tfidf_dense = tfidf_matrix.toarray()
            sim_matrix = cos_sim(tfidf_dense)

        n = len(texts)
        # Normalise CE scores to [0, 1] for relevance weighting
        min_s, max_s = min(ce_scores), max(ce_scores)
        score_range = max_s - min_s or 1.0
        norm_scores = [(s - min_s) / score_range for s in ce_scores]

        # Start with the highest-relevance candidate (index 0 — already CE-sorted)
        selected_indices = [0]
        remaining = list(range(1, n))

        while len(selected_indices) < top_k and remaining:
            best_idx = None
            best_score = float("-inf")

            for i in remaining:
                relevance = norm_scores[i]
                redundancy = float(np.max([sim_matrix[i][j] for j in selected_indices]))
                mmr_score = lambda_val * relevance - (1 - lambda_val) * redundancy
                if mmr_score > best_score:
                    best_score = mmr_score
                    best_idx = i

            if best_idx is not None:
                selected_indices.append(best_idx)
                remaining.remove(best_idx)

        return [(candidates[i][0], candidates[i][1], candidates[i][2]) for i in selected_indices]

    except Exception:
        return [(doc, meta, doc_id) for doc, meta, doc_id, _ in candidates[:top_k]]


def _truncate_to_budget(
    ranked: list[tuple[str, dict[str, Any], str]],
    max_tokens: int,
    return_details: bool = False,
) -> Any:
    """Fit chunks into the context token budget, truncating oversized ones.

    Uses tiktoken (cl100k_base) for accurate token counting when available (B4);
    falls back to a 4-chars-per-token heuristic.

    Instead of skipping chunks that exceed the remaining budget, truncates them
    to fit — preserving at least partial content from highly-ranked chunks.
    Set RAG_MAX_CONTEXT_TOKENS to control the budget (default 6000 tokens).
    """
    try:
        enc = _get_tokenizer()

        def _count_tokens(text: str) -> int:
            return max(1, len(enc.encode(text)))

        def _truncate_text(text: str, max_tok: int) -> str:
            tokens = enc.encode(text)
            if len(tokens) <= max_tok:
                return text
            truncated = enc.decode(tokens[:max_tok])
            # Try to end at a sentence boundary
            last_period = truncated.rfind(". ")
            if last_period > len(truncated) // 2:
                return truncated[:last_period + 1] + " [truncated]"
            return truncated + "… [truncated]"
    except Exception:
        def _count_tokens(text: str) -> int:  # type: ignore[misc]
            return max(1, len(text) // 4)

        def _truncate_text(text: str, max_tok: int) -> str:  # type: ignore[misc]
            max_chars = max_tok * 4
            if len(text) <= max_chars:
                return text
            truncated = text[:max_chars]
            last_period = truncated.rfind(". ")
            if last_period > len(truncated) // 2:
                return truncated[:last_period + 1] + " [truncated]"
            return truncated + "… [truncated]"

    budget = max_tokens
    kept: list[tuple[str, dict[str, Any], str]] = []
    truncated_chunks: list[dict[str, Any]] = []
    skipped_chunks: list[dict[str, Any]] = []
    for chunk, meta, id_ in ranked:
        chunk_tokens = _count_tokens(chunk)
        if chunk_tokens <= budget:
            kept.append((chunk, meta, id_))
            budget -= chunk_tokens
        elif budget >= 50:
            # Truncate to fit remaining budget instead of skipping
            truncated = _truncate_text(chunk, budget)
            kept.append((truncated, meta, id_))
            truncated_chunks.append({
                "id": id_,
                "original_tokens": chunk_tokens,
                "truncated_to": _count_tokens(truncated),
                "source": (meta or {}).get("source", "unknown"),
            })
            budget -= _count_tokens(truncated)
        else:
            skipped_chunks.append({
                "id": id_,
                "tokens": chunk_tokens,
                "source": (meta or {}).get("source", "unknown"),
            })
        if budget <= 0:
            break
    diagnostics = {
        "used_tokens": max_tokens - budget,
        "remaining_tokens": budget,
        "truncated_chunks": truncated_chunks,
        "truncated_count": len(truncated_chunks),
        "skipped_chunks": skipped_chunks,
        "skipped_count": len(skipped_chunks),
    }
    return (kept, diagnostics) if return_details else kept


def get_last_rag_query_diagnostics() -> dict[str, Any]:
    """Return the latest retrieval diagnostics for UI and eval reporting."""
    return dict(_last_rag_query_diagnostics)


_cross_encoder_cache: Any = None
_cross_encoder_model_name: str = ""


def _get_cross_encoder():
    """Load and cache CrossEncoder, rebuilding if RAG_CROSS_ENCODER_MODEL changes (B6)."""
    global _cross_encoder_cache, _cross_encoder_model_name
    model_name = os.getenv("RAG_CROSS_ENCODER_MODEL", "cross-encoder/ms-marco-MiniLM-L-6-v2")
    if _cross_encoder_cache is None or model_name != _cross_encoder_model_name:
        from sentence_transformers import CrossEncoder

        os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")
        os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
        try:
            from transformers.utils import logging as hf_logging

            hf_logging.set_verbosity_error()
        except Exception:
            pass
        _cross_encoder_cache = CrossEncoder(model_name)
        _cross_encoder_model_name = model_name
    return _cross_encoder_cache


@lru_cache(maxsize=1)
def _get_tokenizer():
    """Load and cache a tiktoken encoder for accurate token counting (B4).

    Uses cl100k_base (GPT-4/3.5-turbo tokeniser) as a good general approximation.
    """
    import tiktoken

    return tiktoken.get_encoding("cl100k_base")


@tool(parse_docstring=True)
def ingest_rag_documents(
    file_paths: str,
    project: str = "Default",
    theme: str = "",
    chunking_method: str = "semantic",
    chunk_size: int = 1500,
    chunk_overlap: int = 300,
    breakpoint_threshold: float = 95.0,
) -> str:
    """Ingest documents into the local RAG vector store.

    Args:
        file_paths: Comma/newline separated file paths. Relative paths are resolved from project root.
        project: Top-level project name for organising documents (e.g. 'Climate Policy Review 2026').
        theme: Sub-category within the project (e.g. 'Mitigation', 'Adaptation').
        chunking_method: Splitting strategy. Options: semantic (default, topic-boundary via Ollama embeddings), recursive (character-count, faster), contextual (semantic + LLM context per chunk, slow), late_chunking (semantic + neighbouring-chunk context window), semantic_contextual_late (all three combined, best quality, slowest).
        chunk_size: Max chars per chunk when chunking_method='recursive'.
        chunk_overlap: Overlap chars between chunks when chunking_method='recursive'.
        breakpoint_threshold: Percentile threshold (0-100) for SemanticChunker break detection.
            Higher = fewer, larger chunks. Lower = more, smaller chunks. Default 95.

    Returns:
        Ingestion summary including successes/failures, chunk counts, and method used.
    """
    paths = _normalize_file_list(file_paths)
    try:
        result = ingest_rag_paths(
            paths=paths,
            project=project,
            theme=theme,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            chunking_method=chunking_method,
            breakpoint_threshold=breakpoint_threshold,
        )
    except Exception as e:
        return f"[ERROR] Failed ingestion: {e}"

    lines = [
        f"[OK] Ingestion complete. Loaded files: {result['loaded_files']}/{result['total_files']}",
        f"[OK] Total chunks upserted: {result['added_chunks']}",
        f"[OK] Project: {result['project']} | Theme: {result['theme']}",
        f"[OK] Chunking method: {result.get('method_used', 'unknown')}",
    ]
    if result["failures"]:
        lines.append("[WARN] Failures:")
        lines.extend([f"- {f}" for f in result["failures"]])
    return "\n".join(lines)


@tool(parse_docstring=True)
def ingest_web_search_results(
    query: str,
    project: str = "Default",
    theme: str = "",
    max_results: int = 3,
    topic: str = "general",
    chunking_method: str = "recursive",
    chunk_size: int = 1500,
    chunk_overlap: int = 300,
    breakpoint_threshold: float = 95.0,
) -> str:
    """Search the web and ingest the results into the local RAG vector store.

    Args:
        query: Search query to execute.
        project: Top-level project name for organising documents.
        theme: Sub-category within the project.
        max_results: Maximum number of web results to ingest.
        topic: Tavily topic filter - 'general', 'news', or 'finance'.
        chunking_method: Splitting strategy for the fetched web content.
        chunk_size: Max chars per chunk when chunking_method='recursive'.
        chunk_overlap: Overlap chars between chunks when chunking_method='recursive'.
        breakpoint_threshold: Percentile threshold (0-100) for SemanticChunker break detection.

    Returns:
        Ingestion summary including successes/failures, chunk counts, and method used.
    """
    try:
        from research_agent.tools import TAVILY_AVAILABLE, fetch_webpage_content, tavily_client
    except Exception as exc:
        return f"[ERROR] Research tools unavailable: {exc}"

    if not TAVILY_AVAILABLE or tavily_client is None:
        return "[WARN] Web search is not available. TAVILY_API_KEY is not set."

    try:
        search_results = tavily_client.search(query, max_results=max_results, topic=topic)
    except Exception as exc:
        return f"[ERROR] Web search failed: {exc}"

    records: list[dict[str, Any]] = []
    for index, result in enumerate(search_results.get("results", []), start=1):
        url = str(result.get("url", "") or "").strip()
        title = str(result.get("title", "") or url or f"web-result-{index}").strip()
        if not url:
            continue
        webpage_content = fetch_webpage_content(url)
        fallback_content = str(result.get("content", "") or "").strip()
        body = webpage_content.strip()
        if not body or body.lower().startswith("error fetching content"):
            body = fallback_content
        if not body:
            continue

        domain = ""
        try:
            domain = httpx.URL(url).host or ""
        except Exception:
            domain = ""

        web_text = "\n".join(
            part for part in [
                f"Title: {title}",
                f"URL: {url}",
                f"Domain: {domain}" if domain else "",
                f"Search query: {query}",
                body,
            ] if part
        )
        records.append(
            {
                "source": title,
                "text": web_text,
                "source_fingerprint": _fingerprint_text(f"web:{url}"),
                "metadata": {
                    "source_type": "web_search",
                    "web_url": url,
                    "web_title": title,
                    "web_query": query,
                    "web_topic": topic,
                    "web_domain": domain,
                    "web_snippet": fallback_content[:500],
                    "doc_title": title[:300],
                    "doc_authors": domain[:300],
                    "doc_year": str(result.get("published_date", "") or "")[:20],
                },
            }
        )

    if not records:
        return f"[WARN] No web results could be fetched for '{query}'."

    try:
        result = _ingest_document_records(
            records,
            project=project,
            theme=theme,
            chunking_method=chunking_method,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            breakpoint_threshold=breakpoint_threshold,
        )
    except Exception as exc:
        return f"[ERROR] Failed ingestion: {exc}"

    lines = [
        f"[OK] Web search ingestion complete for query: {query}",
        f"[OK] Loaded results: {result['loaded_files']}/{result['total_files']}",
        f"[OK] Total chunks upserted: {result['added_chunks']}",
        f"[OK] Project: {result['project']} | Theme: {result['theme']}",
        f"[OK] Chunking method: {result.get('method_used', 'unknown')}",
    ]
    if result["failures"]:
        lines.append("[WARN] Failures:")
        lines.extend([f"- {f}" for f in result["failures"]])
    return "\n".join(lines)


@tool(parse_docstring=True)
def list_rag_documents(project: str = "Default") -> str:
    """List indexed documents and themes from the RAG store.

    Args:
        project: Project name used as the RAG partition boundary.

    Returns:
        Document summary grouped by source and theme.
    """
    project_name = _normalize_project(project)
    summary = get_rag_index_summary(project=project_name)
    count = summary["total_chunks"]
    if count == 0:
        return f"No documents indexed in RAG store for project '{project_name}'."

    lines = [f"Indexed RAG documents for project '{project_name}':"]
    for source, info in sorted(summary["files"].items()):
        themes = ", ".join(sorted(info["themes"]))
        lines.append(f"- {source}: {info['chunks']} chunks | themes: {themes}")
    lines.append(f"Total chunks: {count}")
    return "\n".join(lines)


@tool(parse_docstring=True)
def clear_rag_documents(
    project: str = "Default",
    source: str = "",
    theme: str = "",
) -> str:
    """Delete documents from RAG store by source and/or theme.

    Args:
        project: Project name used as the deletion boundary.
        source: Optional source filename to delete.
        theme: Optional theme name to delete.

    Returns:
        Deletion summary.
    """
    project_name = _normalize_project(project)
    deleted = delete_rag_documents(
        project=project_name,
        source=source,
        theme=theme,
        allow_delete_all=not (source or theme),
    )
    if deleted == 0:
        return (
            f"No matching chunks found for deletion in project '{project_name}'."
        )
    return (
        f"[OK] Deleted {deleted} chunks from RAG store for project "
        f"'{project_name}'."
    )


@tool(parse_docstring=True)
def rag_retrieve(
    query: str,
    top_k: int = 5,
    mode: str = "Top-K Globally",
    max_files: int = 5,
    fetch_k: int = 100,
    project: str = "",
    themes: str = "",
    modalities: str = "",
    use_cache: bool = True,
) -> str:
    """Retrieve and rerank chunks from RAG vector store.

    Args:
        query: User query for retrieval.
        top_k: Number of chunks to keep after reranking. Default 5.
               Use 8-10 with 'Top-K Per File' for synthesis/presentation tasks.
        mode: Ranking mode:
              'Top-K Globally' — best chunks globally (focused Q&A, academic).
              'Top-K Per File' — best chunks per document (synthesis, presentation).
              'MMR'            — Maximal Marginal Relevance, removes near-duplicates
                                 (best for synthesis with diverse sources).
              'Hybrid'         — dense vector + BM25 keyword search fused via RRF,
                                 then CrossEncoder reranked (best recall for exact terms).
        max_files: Max files considered when mode is 'Top-K Per File' or 'MMR'.
        fetch_k: Number of initial chunks to fetch from vector DB before reranking.
        project: Optional project name filter (exact match). Takes priority over themes.
        themes: Optional comma-separated themes filter (within the project if set).
        modalities: Optional comma-separated modality filter (`text`, `table`, `image`).

    Returns:
        Formatted retrieved context with citation tokens.
    """
    project_name = _normalize_project(project) if project else ""
    theme_list = _parse_themes(themes)
    modality_list = _parse_modalities(modalities)
    diagnostics: dict[str, Any] = {
        "query": query,
        "project": project_name,
        "themes": theme_list,
        "modalities": modality_list,
        "mode": mode,
        "top_k": top_k,
        "fetch_k": fetch_k,
        "max_files": max_files,
    }

    # --- Result cache check ---
    cache_key = _cache_key(query, project, themes, mode, top_k, fetch_k)
    if use_cache:
        cached = _cache_get(cache_key)
        if cached is not None:
            return cached

    def _store_diagnostics(**updates: Any) -> None:
        _last_rag_query_diagnostics.clear()
        _last_rag_query_diagnostics.update(diagnostics)
        _last_rag_query_diagnostics.update(updates)

    # Build metadata where filter (project + themes + modalities)
    filters: list[dict[str, Any]] = []
    if project_name:
        filters.append({"project": {"$eq": project_name}})
    if theme_list:
        filters.append({"theme": {"$in": theme_list}})
    if modality_list:
        filters.append({"modality": {"$in": modality_list}})

    if len(filters) > 1:
        where_filter: dict | None = {
            "$and": filters
        }
    elif filters:
        where_filter = filters[0]
    else:
        where_filter = None

    # --- HyDE: Hypothetical Document Embedding ---
    # Generate a hypothetical answer to use as a richer query embedding.
    # Enabled via RAG_ENABLE_HYDE=true (default: auto — enabled only for Hybrid/MMR).
    hyde_setting = os.getenv("RAG_ENABLE_HYDE", "auto").lower()
    use_hyde = hyde_setting == "true" or (hyde_setting == "auto" and mode in ("Hybrid", "MMR"))
    embedding_query = query
    if use_hyde:
        hyde_doc = _generate_hyde_document(query)
        if hyde_doc:
            embedding_query = hyde_doc
            diagnostics["hyde_used"] = True
            diagnostics["hyde_length"] = len(hyde_doc)
        else:
            diagnostics["hyde_used"] = False

    # --- Query Decomposition ---
    # For complex queries, decompose into sub-queries and merge results via RRF.
    # Enabled via RAG_ENABLE_DECOMPOSITION=true (default: auto — for MMR/Hybrid).
    decomp_setting = os.getenv("RAG_ENABLE_DECOMPOSITION", "auto").lower()
    use_decomp = decomp_setting == "true" or (decomp_setting == "auto" and mode in ("MMR", "Hybrid", "Top-K Per File"))
    sub_queries = [embedding_query]
    if use_decomp:
        sub_queries = _decompose_query(query)
        # If decomposed, apply HyDE to original query only (already done above)
        if len(sub_queries) > 1 and use_hyde and embedding_query != query:
            sub_queries[0] = embedding_query  # Replace original with HyDE version
        diagnostics["sub_queries"] = sub_queries
        diagnostics["decomposition_used"] = len(sub_queries) > 1

    # Check that at least one collection has data before querying
    all_collections = _list_rag_collections()
    if not all_collections or all(c.count() == 0 for c in all_collections):
        _store_diagnostics(status="empty-store", candidate_count=0)
        return "[WARN] RAG store is empty. Ingest documents first."

    # Query across ALL matching collections — handles migration from old collection names
    # When query decomposition is active, run each sub-query and fuse via RRF.
    if len(sub_queries) > 1:
        sub_results: list[list[tuple[str, dict[str, Any], str]]] = []
        for sq in sub_queries:
            sq_docs, sq_metas, sq_ids = _query_all_matching_collections(
                query=sq,
                n_results=fetch_k,
                where_filter=where_filter,
            )
            sub_results.append(list(zip(sq_docs, sq_metas, sq_ids)))
        fused_decomp = _reciprocal_rank_fusion(sub_results)
        docs = [d for d, _, _ in fused_decomp]
        metas = [m for _, m, _ in fused_decomp]
        ids = [i for _, _, i in fused_decomp]
    else:
        docs, metas, ids = _query_all_matching_collections(
            query=sub_queries[0],
            n_results=fetch_k,
            where_filter=where_filter,
        )
    diagnostics["candidate_count"] = len(docs)

    if not docs:
        _store_diagnostics(status="no-match", candidate_count=0)
        return "[WARN] No relevant chunks found for query."

    # B1: Hybrid mode — fuse dense (cross-collection) + BM25 results via RRF,
    # then pass the merged pool to CrossEncoder reranking.
    # Uses cached BM25 index for performance.
    if mode == "Hybrid":
        try:
            # Use cached BM25 index instead of fetching all docs from scratch
            bm25_pool_docs: list[str] = []
            bm25_pool_metas: list[dict[str, Any]] = []
            bm25_pool_ids: list[str] = []
            for coll_obj in all_collections:
                coll_name = coll_obj.name if hasattr(coll_obj, "name") else str(coll_obj)
                bm25_index = _get_bm25_index(coll_name, coll_obj)
                if bm25_index is None:
                    continue
                _, _, pool_docs, pool_metas, pool_ids = bm25_index
                bm25_pool_docs.extend(pool_docs)
                bm25_pool_metas.extend(pool_metas)
                bm25_pool_ids.extend(pool_ids)
            if where_filter and bm25_pool_docs:
                filtered = [
                    (d, m, i)
                    for d, m, i in zip(bm25_pool_docs, bm25_pool_metas, bm25_pool_ids)
                    if _meta_matches_filter(m, where_filter)
                ]
                if filtered:
                    bm25_pool_docs = [t[0] for t in filtered]
                    bm25_pool_metas = [t[1] for t in filtered]
                    bm25_pool_ids = [t[2] for t in filtered]
                else:
                    bm25_pool_docs, bm25_pool_metas, bm25_pool_ids = [], [], []
            # Use prebuilt BM25 from cache for each collection, or fall back to on-the-fly
            bm25_results = _bm25_search(
                query=query,
                documents=bm25_pool_docs,
                metadatas=bm25_pool_metas,
                ids=bm25_pool_ids,
                top_k=fetch_k,
            )
        except Exception:
            bm25_results = []
        dense_results = [(d, m, i) for d, m, i in zip(docs, metas, ids)]
        fused = _reciprocal_rank_fusion([dense_results, bm25_results])
        diagnostics.update(
            {
                "dense_candidate_count": len(dense_results),
                "bm25_pool_count": len(bm25_pool_docs),
                "bm25_result_count": len(bm25_results),
                "fused_candidate_count": len(fused),
            }
        )
        docs = [d for d, _, _ in fused]
        metas = [m for _, m, _ in fused]
        ids = [i for _, _, i in fused]
        rerank_mode = "Top-K Globally"
    else:
        rerank_mode = mode

    min_score = float(os.getenv("RAG_MIN_RERANK_SCORE", "0.0"))
    try:
        rerank_result = _rerank_chunks(
            query=query,
            documents=[str(d) for d in docs],
            metadatas=[m or {} for m in metas],
            ids=[str(i) for i in ids],
            top_k=top_k,
            mode=rerank_mode,
            max_files=max_files,
            min_score=min_score,
            return_details=True,
        )
    except TypeError:
        rerank_result = _rerank_chunks(
            query=query,
            documents=[str(d) for d in docs],
            metadatas=[m or {} for m in metas],
            ids=[str(i) for i in ids],
            top_k=top_k,
            mode=rerank_mode,
            max_files=max_files,
            min_score=min_score,
        )

    if isinstance(rerank_result, tuple) and len(rerank_result) == 2:
        ranked, rerank_details = rerank_result
    else:
        ranked = rerank_result
        rerank_details = {
            "rerank_count": len(ranked),
            "filtered_count": len(ranked),
            "selected_files": [],
            "file_scores": {},
        }
    diagnostics.update(rerank_details)

    if not ranked:
        _store_diagnostics(
            status="below-threshold",
            final_chunk_ids=[],
            final_chunk_count=0,
        )
        return (
            "[WARN] All retrieved chunks fell below the relevance score threshold "
            f"(RAG_MIN_RERANK_SCORE={min_score}). Try lowering the threshold or rephrasing the query."
        )

    # Token budget enforcement — prevents silent context window overflow
    max_tokens = int(os.getenv("RAG_MAX_CONTEXT_TOKENS", "6000"))
    try:
        budget_result = _truncate_to_budget(
            ranked,
            max_tokens=max_tokens,
            return_details=True,
        )
    except TypeError:
        budget_result = _truncate_to_budget(
            ranked,
            max_tokens=max_tokens,
        )

    if isinstance(budget_result, tuple) and len(budget_result) == 2:
        ranked, budget_details = budget_result
    else:
        ranked = budget_result
        budget_details = {
            "used_tokens": None,
            "remaining_tokens": None,
            "skipped_chunks": [],
            "skipped_count": 0,
        }
    diagnostics.update(
        {
            "max_context_tokens": max_tokens,
            **budget_details,
            "final_chunk_ids": [chunk_id for _chunk, _meta, chunk_id in ranked],
            "final_chunk_count": len(ranked),
        }
    )
    selected_modalities: dict[str, int] = {}
    selected_chunking_methods: dict[str, int] = {}
    selected_table_extraction_methods: dict[str, int] = {}
    for _chunk, meta, _chunk_id in ranked:
        modality = str((meta or {}).get("modality", "text") or "text").lower()
        selected_modalities[modality] = selected_modalities.get(modality, 0) + 1
        chunking_method = str((meta or {}).get("chunking_method", "") or "").strip().lower()
        if chunking_method:
            selected_chunking_methods[chunking_method] = selected_chunking_methods.get(chunking_method, 0) + 1
        table_method = str((meta or {}).get("table_extraction_method", "") or "").strip().lower()
        if table_method:
            selected_table_extraction_methods[table_method] = (
                selected_table_extraction_methods.get(table_method, 0) + 1
            )
    diagnostics["selected_modalities"] = selected_modalities
    diagnostics["selected_chunking_methods"] = selected_chunking_methods
    diagnostics["selected_table_extraction_methods"] = selected_table_extraction_methods

    if not ranked:
        _store_diagnostics(status="budget-truncated")
        return (
            "[WARN] Retrieved chunks exceeded the active context token budget. "
            "Increase RAG_MAX_CONTEXT_TOKENS or reduce top_k."
        )

    _store_diagnostics(status="ok")

    n_chunks = len(ranked)
    valid_ids_str = ", ".join(f"[R{i}]" for i in range(1, n_chunks + 1))
    lines = [
        f"Retrieved {n_chunks} chunk(s) for: _{query}_",
        # Explicit guard: the answering LLM must only cite IDs that exist here.
        f"VALID CITATION IDs: {valid_ids_str}. "
        "Only cite these IDs — do not invent, reuse, or extrapolate additional IDs.",
        "",
    ]
    references: list[tuple[str, str, str]] = []

    for idx, (chunk, meta, chunk_id) in enumerate(ranked, start=1):
        source = meta.get("source", "unknown")
        section = meta.get("section", "")
        topics = meta.get("topics", "")
        page_number = meta.get("page_number", "")
        modality = meta.get("modality", "text") or "text"
        asset_path = meta.get("asset_path", "")
        ref_id = f"R{idx}"
        apa_ref, year = _build_apa_reference(meta)
        intext_author = _build_intext_author(meta)
        intext_citation = f"({intext_author}, {year}{',' + ' p. ' + str(page_number) if page_number else ''})"
        references.append((ref_id, apa_ref, intext_citation))

        # Compact chunk header
        header = f"[{ref_id}] {source}" if modality == "text" else f"[{ref_id}] {modality} | {source}"
        if page_number:
            header += f" p.{page_number}"
        if section:
            header += f" § {section}"
        lines.append(header)
        if topics:
            lines.append(f"  Topics: {topics}")
        if asset_path:
            lines.append(f"  Asset: {asset_path}")

        # Prefer the verbatim source text stored at ingestion time over the
        # embedding-enriched text which may contain LLM-generated [Context: ...]
        # or [LateCtx: ...] prefixes — those are for retrieval only, not quotable.
        display_text = meta.get("chunk_source_text") or str(chunk)
        lines.append(display_text)
        lines.append("")

    lines.append("---")
    lines.append(
        f"Only cite the {n_chunks} IDs listed above ({valid_ids_str}). "
        "Copy APA reference lines verbatim from the References section below. "
        "Do not invent authors, titles, page numbers, or any other bibliographic detail."
    )
    lines.append("References:")
    for ref_id, apa_ref, intext in references:
        lines.append(f"  [{ref_id}] {apa_ref}  {intext}")
    result = "\n".join(lines)
    if use_cache:
        _cache_put(cache_key, result)
    return result


@tool(parse_docstring=True)
def rag_think_tool(reflection: str) -> str:
    """Record a short reflection before/after RAG steps.

    Args:
        reflection: Reflection text.

    Returns:
        Confirmation string.
    """
    return f"RAG reflection recorded: {reflection}"
